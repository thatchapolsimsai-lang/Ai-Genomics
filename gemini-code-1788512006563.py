import gc
import base64
import hashlib
import json
import os
import re
import requests
import shutil
import time
from io import BytesIO, StringIO
from pathlib import Path
from typing import List
from urllib.error import HTTPError, URLError

import streamlit as st
from Bio import Entrez, SeqIO
from Bio.Seq import Seq
from docx import Document as WordDocument
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_core.messages import HumanMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

try:
    from langchain_groq import ChatGroq
except ImportError:
    ChatGroq = None

try:
    from langchain_nvidia_ai_endpoints import ChatNVIDIA
except ImportError:
    ChatNVIDIA = None

# =============================================================================
# Configuration and shared helpers
# =============================================================================

PROJECT_FOLDER = Path(__file__).resolve().parent
PDF_FOLDER = PROJECT_FOLDER
VECTOR_DB_PATH = PROJECT_FOLDER / "vector_db"
FINGERPRINT_FILE = VECTOR_DB_PATH / "pdf_fingerprint.txt"

RAG_MODEL = "gemini-3.6-flash"
EMBEDDING_MODEL = "gemini-embedding-2-preview"
EMBED_BATCH_SIZE = 25
EMBED_BATCH_DELAY_SECONDS = 2.0
EMBED_MAX_RETRIES = 6
RETRIEVER_K = 5
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
ONLINE_RESULT_LIMIT = 5
ONLINE_TIMEOUT_SECONDS = 20
UNIPROT_ENTRY_ENDPOINT = "https://rest.uniprot.org/uniprotkb/{accession_id}?format=json"
UNIPROT_GENE_ENDPOINT = "https://rest.uniprot.org/uniprotkb/search?query=gene:{gene_name}&format=json"
UNIPROT_ORGANISM_ENDPOINT = "https://rest.uniprot.org/uniprotkb/search?query=organism_id:{tax_id}&format=json"
API_TIMEOUT_SECONDS = 20

MODEL_OPTIONS = {
    "Gemini 3.6 Flash (Google)": {
        "provider": "google",
        "model": "gemini-3.6-flash",
        "secret": "GOOGLE_API_KEY",
    },
    "Claude 3.5 (9arm API)": {
        "provider": "9arm",
        "model": "claude-3.5-sonnet",
        "secret": "NINEARM_API_KEY",
    },
    "Llama 3 8B (Groq)": {
        "provider": "groq",
        "model": "llama3-8b-8192",
        "secret": "GROQ_API_KEY",
    },
    "Mixtral 8x7B (Groq)": {
        "provider": "groq",
        "model": "mixtral-8x7b-32768",
        "secret": "GROQ_API_KEY",
    },
    "Llama 3 70B (NVIDIA NIM)": {
        "provider": "nvidia",
        "model": "meta/llama3-70b-instruct",
        "secret": "NVIDIA_API_KEY",
    },
}

st.set_page_config(page_title="AI Research Workbench", page_icon="🧬", layout="wide")


def get_api_key() -> str:
    key = os.getenv("GOOGLE_API_KEY", "").strip()
    if key:
        return key
    try:
        return str(st.secrets.get("GOOGLE_API_KEY", "")).strip()
    except Exception:
        return ""


GOOGLE_API_KEY = get_api_key()
if GOOGLE_API_KEY:
    os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY


def is_rate_limit_error(exc: Exception) -> bool:
    text = str(exc).upper()
    return any(item in text for item in ("429", "RESOURCE_EXHAUSTED", "QUOTA", "RATE LIMIT"))


def extract_text(response) -> str:
    content = getattr(response, "content", response)
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        return "\n".join(
            block if isinstance(block, str) else block.get("text", "")
            for block in content
            if isinstance(block, str) or isinstance(block, dict)
        ).strip()
    if isinstance(content, dict):
        return str(content.get("text", "")).strip()
    return ""


def format_bytes(size: int) -> str:
    if size < 1024:
        return f"{size} B"
    if size < 1024 ** 2:
        return f"{size / 1024:.1f} KB"
    return f"{size / (1024 ** 2):.1f} MB"


def get_configured_key(key_name: str) -> str:
    key = os.getenv(key_name, "").strip()
    if key:
        return key
    try:
        return str(st.secrets.get(key_name, "")).strip()
    except Exception:
        return ""

# ตั้งค่า NCBI API Key ส่วนกลาง
NCBI_API_KEY = get_configured_key("NCBI_API_KEY")
if NCBI_API_KEY:
    Entrez.api_key = NCBI_API_KEY


def get_active_llm(model_choice: str, api_key: str):
    config = MODEL_OPTIONS.get(model_choice)
    if config is None:
        raise ValueError("ไม่พบโมเดลที่เลือก")
    if not api_key:
        raise ValueError(f"ยังไม่ได้ตั้งค่า API key สำหรับ {model_choice}")

    if config["provider"] == "google":
        return ChatGoogleGenerativeAI(model=config["model"], temperature=0.2, google_api_key=api_key)
    if config["provider"] == "9arm":
        try:
            from langchain_openai import ChatOpenAI
        except ImportError:
            raise RuntimeError("กรุณาติดตั้ง langchain-openai ก่อนใช้ 9arm API")
        return ChatOpenAI(
            model=config["model"],
            temperature=0.2,
            api_key=api_key,
            base_url="https://api.9arm.com/v1" # ปรับ URL ตาม Endpoint จริงของ 9arm
        )
    if config["provider"] == "groq":
        if ChatGroq is None:
            raise RuntimeError("กรุณาติดตั้ง langchain-groq ก่อนใช้โมเดล Groq")
        return ChatGroq(model=config["model"], temperature=0.2, groq_api_key=api_key)
    if ChatNVIDIA is None:
        raise RuntimeError("กรุณาติดตั้ง langchain-nvidia-ai-endpoints ก่อนใช้ NVIDIA NIM")
    return ChatNVIDIA(model=config["model"], temperature=0.2, api_key=api_key)


# =============================================================================
# Module 1: Local Document RAG
# =============================================================================

def get_pdf_files() -> List[Path]:
    return sorted(
        (path for path in PDF_FOLDER.iterdir() if path.is_file() and path.suffix.lower() == ".pdf"),
        key=lambda path: path.name.lower(),
    )


def calculate_pdf_fingerprint(pdf_files: List[Path]) -> str:
    digest = hashlib.sha256()
    digest.update(f"embedding:{EMBEDDING_MODEL}|chunk:{CHUNK_SIZE}|overlap:{CHUNK_OVERLAP}".encode())
    for pdf_file in pdf_files:
        stat = pdf_file.stat()
        digest.update(f"{pdf_file.name}:{stat.st_size}:{stat.st_mtime_ns}".encode())
    return digest.hexdigest()


def database_is_current(pdf_files: List[Path]) -> bool:
    if not VECTOR_DB_PATH.exists() or not FINGERPRINT_FILE.exists():
        return False
    try:
        return FINGERPRINT_FILE.read_text(encoding="utf-8").strip() == calculate_pdf_fingerprint(pdf_files)
    except OSError:
        return False


def remove_vector_database(max_attempts: int = 8) -> None:
    if not VECTOR_DB_PATH.exists():
        return
    gc.collect()
    last_error = None
    for attempt in range(max_attempts):
        try:
            shutil.rmtree(VECTOR_DB_PATH)
            return
        except PermissionError as exc:
            last_error = exc
            gc.collect()
            if attempt < max_attempts - 1:
                time.sleep(min(2.0, 0.25 * (attempt + 1)))
    raise PermissionError("vector_db ยังถูกใช้งานอยู่ กรุณาปิด Streamlit instance อื่นแล้วลองใหม่") from last_error


def load_pdf_documents(pdf_files: List[Path]):
    documents, errors = [], []
    for pdf_file in pdf_files:
        try:
            reader = PdfReader(str(pdf_file), strict=False)
            found_text = False
            for page_number, page in enumerate(reader.pages, start=1):
                try:
                    text = (page.extract_text() or "").strip()
                    if text:
                        found_text = True
                        documents.append(Document(page_content=text, metadata={"source": pdf_file.name, "page": page_number}))
                except Exception as exc:
                    errors.append(f"{pdf_file.name} หน้า {page_number}: {exc}")
            if not found_text:
                errors.append(f"{pdf_file.name}: ไม่พบข้อความ")
        except Exception as exc:
            errors.append(f"{pdf_file.name}: {exc}")
    return documents, errors


class RateLimitedGeminiEmbeddings(Embeddings):
    def __init__(self, progress_callback=None):
        self.inner = GoogleGenerativeAIEmbeddings(model=EMBEDDING_MODEL)
        self.progress_callback = progress_callback

    def _embed_with_retry(self, texts):
        last_error = None
        for attempt in range(EMBED_MAX_RETRIES):
            try:
                return self.inner.embed_documents(texts)
            except Exception as exc:
                last_error = exc
                if not is_rate_limit_error(exc):
                    raise
                delay = min(120.0, max(30.0, EMBED_BATCH_DELAY_SECONDS * (2 ** min(attempt, 4))))
                if self.progress_callback:
                    self.progress_callback(f"Gemini quota: รอ {delay:.0f} วินาที")
                time.sleep(delay)
        raise RuntimeError(f"สร้าง embedding ไม่สำเร็จ: {last_error}")

    def embed_documents(self, texts: List[str]):
        results = []
        for start in range(0, len(texts), EMBED_BATCH_SIZE):
            batch = texts[start:start + EMBED_BATCH_SIZE]
            results.extend(self._embed_with_retry(batch))
            if self.progress_callback:
                self.progress_callback(f"Embedding {min(start + len(batch), len(texts))}/{len(texts)} chunks")
            if start + len(batch) < len(texts):
                time.sleep(EMBED_BATCH_DELAY_SECONDS)
        return results

    def embed_query(self, text: str):
        return self._embed_with_retry([text])[0]


def build_vector_database(pdf_files, embeddings, progress_callback=None):
    raw_documents, errors = load_pdf_documents(pdf_files)
    if not raw_documents:
        return None, "ไม่สามารถอ่านข้อความจาก PDF ได้\n" + "\n".join(errors)
    splits = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP).split_documents(raw_documents)
    if not splits:
        return None, "ไม่พบข้อความที่แบ่งเป็น chunks ได้"
    if progress_callback:
        progress_callback(f"อ่านได้ {len(raw_documents)} หน้า, แบ่งเป็น {len(splits)} chunks")
    remove_vector_database()
    vectorstore = Chroma(persist_directory=str(VECTOR_DB_PATH), embedding_function=embeddings, collection_name="mini_rag")
    texts = [doc.page_content for doc in splits]
    ids = [hashlib.sha256(f"{index}:{doc.page_content}".encode()).hexdigest() for index, doc in enumerate(splits)]
    vectorstore._collection.add(ids=ids, documents=texts, metadatas=[doc.metadata for doc in splits], embeddings=embeddings.embed_documents(texts))
    VECTOR_DB_PATH.mkdir(parents=True, exist_ok=True)
    FINGERPRINT_FILE.write_text(calculate_pdf_fingerprint(pdf_files), encoding="utf-8")
    return vectorstore, None


@st.cache_resource
def initialize_rag(pdf_fingerprint: str, model_choice: str, api_key: str):
    pdf_files = get_pdf_files()
    if not GOOGLE_API_KEY:
        return None, None, "ไม่พบ GOOGLE_API_KEY"
    try:
        progress = []
        embeddings = RateLimitedGeminiEmbeddings(progress.append)
        error = None
        if database_is_current(pdf_files):
            vectorstore = Chroma(persist_directory=str(VECTOR_DB_PATH), embedding_function=embeddings, collection_name="mini_rag")
            if vectorstore._collection.count() == 0:
                vectorstore, error = build_vector_database(pdf_files, embeddings, progress.append)
        else:
            vectorstore, error = build_vector_database(pdf_files, embeddings, progress.append)
        if error:
            return None, progress, error
        retriever = vectorstore.as_retriever(search_kwargs={"k": RETRIEVER_K})
        prompt = ChatPromptTemplate.from_template("""
    คุณคือ Expert Data Analyst ที่วิเคราะห์เอกสารอย่างมีหลักฐาน
    ใช้ Context เท่านั้น ห้ามแต่งข้อมูล หากข้อมูลไม่พอให้ระบุไว้ใน uncertainties
    ตอบเป็น JSON object ดิบเท่านั้น ห้ามใช้ Markdown code fence หรือข้อความอื่น
    ต้องตรงกับ schema นี้ทุกครั้ง:
    {{
      "summary": "String",
      "confidence_score": 0,
      "uncertainties": ["String"],
      "next_steps": ["String"]
    }}
    confidence_score ต้องเป็นจำนวนเต็มระหว่าง 0 ถึง 100
    Context:\n{context}\n\nResearch query:\n{question}
    """)
        llm = get_active_llm(model_choice, api_key)
        def format_docs(docs):
            return "\n\n---\n\n".join(doc.page_content for doc in docs) or "ไม่พบเอกสารที่เกี่ยวข้อง"
        chain = ({"context": retriever | format_docs, "question": RunnablePassthrough()} | prompt | llm)
        return vectorstore, (chain, retriever), None
    except Exception as exc:
        return None, None, f"เตรียมฐานข้อมูลไม่สำเร็จ: {exc}"


def fetch_online_open_access_context(query: str):
    context = []
    sources = []

    with st.spinner("กำลังค้นบทความ Open Access จาก Europe PMC..."):
        try:
            response = requests.get(
                "https://www.ebi.ac.uk/europepmc/webservices/rest/search",
                params={
                    "query": f"({query}) AND OPEN_ACCESS:Y",
                    "format": "json",
                    "resultType": "core",
                    "pageSize": ONLINE_RESULT_LIMIT,
                },
                timeout=ONLINE_TIMEOUT_SECONDS,
            )
            response.raise_for_status()
            for item in response.json().get("resultList", {}).get("result", []):
                identifier = f"[PMID: {item.get('pmid')}]" if item.get("pmid") else f"[DOI: {item.get('doi', 'unknown')}]"
                sources.append(identifier)
                context.append({
                    "database": "Europe PMC",
                    "source": identifier,
                    "title": item.get("title", ""),
                    "journal": item.get("journalTitle", ""),
                    "abstract": item.get("abstractText", ""),
                    "year": item.get("pubYear", ""),
                })
        except Exception:
            context.append({"source": "Europe PMC", "status": "data unavailable"})

    with st.spinner("กำลังค้นข้อมูล Open Access จาก OpenAlex..."):
        try:
            response = requests.get(
                "https://api.openalex.org/works",
                params={"search": query, "filter": "is_oa:true", "per-page": ONLINE_RESULT_LIMIT},
                timeout=ONLINE_TIMEOUT_SECONDS,
            )
            response.raise_for_status()
            for item in response.json().get("results", []):
                doi = item.get("doi") or item.get("ids", {}).get("openalex", "unknown")
                identifier = f"[DOI: {doi.replace('https://doi.org/', '')}]" if doi else "[OpenAlex: unknown]"
                sources.append(identifier)
                context.append({
                    "database": "OpenAlex",
                    "source": identifier,
                    "title": item.get("title", ""),
                    "journal": item.get("primary_location", {}).get("source", {}).get("display_name", "") if item.get("primary_location") else "",
                    "year": item.get("publication_year", ""),
                    "open_access": item.get("open_access", {}),
                    "landing_page": item.get("primary_location", {}).get("landing_page_url", "") if item.get("primary_location") else "",
                })
        except Exception:
            context.append({"source": "OpenAlex", "status": "data unavailable"})

    return context, list(dict.fromkeys(sources))


def render_document_rag():
    render_online_research()


def render_online_research():
    st.header("🌐 Open Access Biology Research")
    st.caption("Chatbot ที่ใช้ข้อมูล Open Access ออนไลน์เท่านั้น ไม่ใช้ PDF หรือฐานข้อมูลภายใน")

    # ปรับปรุง UI การอัปโหลดให้ซ่อนตัวเหมือนแอปพลิเคชันแชท
    with st.popover("📎 แนบไฟล์เอกสาร / 📷 ถ่ายภาพ"):
        attached_files = st.file_uploader(
            "อัปโหลดเอกสารหรือรูปภาพ",
            type=["pdf", "txt", "md", "png", "jpg", "jpeg", "webp"],
            accept_multiple_files=True,
            key="online_attachments",
            label_visibility="collapsed"
        )
        camera_image = st.camera_input("ถ่ายภาพ", key="online_camera", label_visibility="collapsed")

    bio_metrics = st.session_state.get("bio_metrics")
    bio_report = st.session_state.get("bio_report", "")
    if bio_metrics:
        st.divider()
        st.subheader("วิเคราะห์ต่อจาก Bioinformatics Agent")
        st.caption("ใช้ผลวิเคราะห์จากหน้า Bio โดยอัตโนมัติ ไม่ต้องคัดลอกข้อมูล")
        bio_handoff_query = st.text_input(
            "คำถามสำหรับผลวิเคราะห์ Bio",
            placeholder="เช่น เปรียบเทียบผล sequence กับบทความ Open Access",
            key="online_bio_handoff_query",
        )
        if st.button("วิเคราะห์ต่อจากผล Bio ครั้งเดียว", key="online_analyze_bio_once"):
            try:
                attachment_content = build_attachment_content(attached_files, camera_image)
                with st.spinner("กำลังค้น Open Access และวิเคราะห์ผลจาก Bioinformatics Agent..."):
                    context, sources = fetch_online_open_access_context(
                        bio_handoff_query or "ชีววิทยาและหน้าที่ของ sequence นี้"
                    )
                    bio_context_prompt = f"""
คุณคือ Advanced Bioinformatics Research Agent
ใช้เฉพาะข้อมูลใน <open_access_context>, ผลวิเคราะห์ Bioinformatics เดิม
และไฟล์แนบของผู้ใช้เท่านั้น ห้ามใช้ความรู้ภายในหรือเดา
แยก Facts และ Inferences และใส่ citation จาก context ต่อท้ายทุก claim

ผลวิเคราะห์ Bioinformatics เดิม:
{bio_report or '(ไม่มีรายงานเดิม)'}
Deterministic metrics:
{json.dumps(bio_metrics, ensure_ascii=False, indent=2)}
<open_access_context>
{json.dumps(context, ensure_ascii=False, indent=2)}
</open_access_context>
คำถาม: {bio_handoff_query or 'วิเคราะห์ผล Bioinformatics นี้ร่วมกับข้อมูล Open Access'}
"""
                    answer = extract_text(
                        get_active_llm(
                            st.session_state.active_model_choice,
                            st.session_state.active_model_key,
                        ).invoke([HumanMessage(content=[bio_context_prompt, *attachment_content[0]])])
                    )
                    st.session_state.online_bio_analysis = answer
                    st.session_state.online_bio_sources = sources
            except Exception:
                st.error("ไม่สามารถวิเคราะห์ต่อจากผล Bioinformatics ได้ กรุณาตรวจสอบ API key หรือการเชื่อมต่อ")
        if st.session_state.get("online_bio_analysis"):
            st.markdown(st.session_state.online_bio_analysis)
            with st.expander("Open Access sources จากการวิเคราะห์ต่อ"):
                for source in st.session_state.get("online_bio_sources", []):
                    st.write(source)

    messages = st.session_state.setdefault("online_chat_messages", [])
    for message in messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            if message.get("sources"):
                with st.expander("Open Access sources"):
                    for source in message["sources"]:
                        st.write(source)

    query = st.chat_input("ถามเกี่ยวกับชีววิทยา ชีวเคมี หรือชีวสารสนเทศ")
    if not query:
        return

    messages.append({"role": "user", "content": query})
    with st.chat_message("user"):
        st.markdown(query)

    with st.chat_message("assistant"):
        try:
            context, sources = fetch_online_open_access_context(query)
            available_records = [
                record for record in context
                if record.get("abstract") or record.get("title")
            ]
            if not available_records:
                answer = "ข้อมูลจากฐานข้อมูล Open Access ที่ดึงมา ณ ขณะนี้ ไม่เพียงพอต่อการวิเคราะห์เชิงลึก"
            else:
                attachment_parts, attachment_names = build_attachment_content(
                    attached_files,
                    camera_image,
                )
                prompt = f"""
คุณคือ Advanced Bioinformatics Research Agent ในรูปแบบ chatbot
ใช้ข้อมูลได้เฉพาะจาก <open_access_context> และไฟล์แนบของผู้ใช้เท่านั้น
ห้ามใช้ pretrained knowledge, memory, assumptions หรือฐานข้อมูลภายใน
ทุก factual claim ต้องมี inline citation จาก identifier ที่ปรากฏใน context ทันที
แยกหัวข้อ Facts (ข้อเท็จจริง) และ Inferences (การอนุมาน) ให้ชัดเจน
Inference ต้องอ้างอิง facts และ citation ที่รองรับเท่านั้น
ห้ามสร้าง citation, PMID, DOI, Accession, PDB หรือผลการทดลองขึ้นเอง
หาก context ไม่พอ ให้ตอบ EXACTLY:
ข้อมูลจากฐานข้อมูล Open Access ที่ดึงมา ณ ขณะนี้ ไม่เพียงพอต่อการวิเคราะห์เชิงลึก
ตอบเป็น Academic Thai และตอบตรงคำถามของผู้ใช้

<open_access_context>
{json.dumps(context, ensure_ascii=False, indent=2)}
</open_access_context>
Attached files: {', '.join(attachment_names) or '(none)'}
User query: {query}
"""
                with st.spinner("กำลังค้นและวิเคราะห์ข้อมูล Open Access..."):
                    response = get_active_llm(
                        st.session_state.active_model_choice,
                        st.session_state.active_model_key,
                    ).invoke([HumanMessage(content=[prompt, *attachment_parts])])
                answer = extract_text(response).strip()
            st.markdown(answer)
            if sources:
                with st.expander("Open Access sources"):
                    for source in sources:
                        st.write(source)
            messages.append({"role": "assistant", "content": answer, "sources": sources})
        except Exception:
            answer = "ไม่สามารถวิเคราะห์ข้อมูล Open Access ได้ กรุณาลองใหม่หรือตรวจสอบ API key และการเชื่อมต่อ"
            st.error(answer)
            messages.append({"role": "assistant", "content": answer, "sources": []})


def parse_analysis_json(raw_response: str):
    cleaned = raw_response.strip()
    if cleaned.startswith("```"):
        raise ValueError("โมเดลส่ง Markdown code fence แทน raw JSON")
    analysis = json.loads(cleaned)
    required = {"summary", "confidence_score", "uncertainties", "next_steps"}
    if not isinstance(analysis, dict) or set(analysis) != required:
        raise ValueError("JSON ไม่ตรง schema ที่กำหนด")
    if not isinstance(analysis["summary"], str):
        raise ValueError("summary ต้องเป็น string")
    if not isinstance(analysis["confidence_score"], int) or isinstance(analysis["confidence_score"], bool):
        raise ValueError("confidence_score ต้องเป็น integer")
    if not 0 <= analysis["confidence_score"] <= 100:
        raise ValueError("confidence_score ต้องอยู่ระหว่าง 0-100")
    for field in ("uncertainties", "next_steps"):
        if not isinstance(analysis[field], list) or not all(isinstance(item, str) for item in analysis[field]):
            raise ValueError(f"{field} ต้องเป็น list ของ string")
    return analysis


# =============================================================================
# Module 2: Biopython deterministic pipeline and Gemini agent
# =============================================================================

def parse_uploaded_sequence(uploaded_file):
    try:
        content = uploaded_file.getvalue().decode("utf-8")
        record = SeqIO.read(StringIO(content), "fasta")
    except Exception as exc:
        raise ValueError(f"รูปแบบ FASTA ไม่ถูกต้อง: {exc}") from exc
    return record.seq, record.id


def parse_raw_sequence(raw_input: str):
    sequence = re.sub(r"\s+", "", raw_input or "").upper()
    if not sequence:
        raise ValueError("กรุณาใส่ nucleotide sequence")
    return Seq(sequence), "raw-sequence"


def fetch_ncbi_sequence(accession: str):
    accession = accession.strip()
    if not accession:
        raise ValueError("กรุณาระบุ NCBI Accession Number")

    Entrez.email = "developer@example.com"
    try:
        with st.spinner("Fetching sequence from GenBank..."):
            with Entrez.efetch(
                db="nucleotide",
                id=accession,
                rettype="fasta",
                retmode="text",
            ) as handle:
                record = SeqIO.read(handle, "fasta")
        return record.seq, record.id
    except HTTPError as exc:
        raise ValueError(f"NCBI ไม่พบ accession หรือปฏิเสธคำขอ ({exc.code})") from exc
    except URLError as exc:
        raise ValueError(f"เชื่อมต่อ NCBI ไม่สำเร็จ: {exc.reason}") from exc
    except Exception as exc:
        raise ValueError(f"ดึง sequence จาก NCBI ไม่สำเร็จ: {exc}") from exc


def calculate_sequence_metrics(sequence: Seq):
    sequence_text = str(sequence).upper()
    invalid = sorted(set(sequence_text) - set("ACGTN"))
    if invalid:
        raise ValueError(f"พบอักขระที่ไม่ใช่ nucleotide: {', '.join(invalid)}")
    gc_count = sequence_text.count("G") + sequence_text.count("C")
    gc_content = gc_count / len(sequence_text) * 100 if sequence_text else 0.0
    mrna = sequence.transcribe()
    protein = mrna.translate(to_stop=True)
    return {"sequence": sequence_text, "length": len(sequence_text), "gc": gc_content, "mrna": str(mrna), "protein": str(protein)}


def resolve_database_query(sequence_id: str, user_query: str, metrics: dict) -> str:
    if sequence_id and sequence_id != "raw-sequence":
        return sequence_id.strip()
    if user_query.strip():
        return user_query.strip()
    return metrics.get("protein", "").strip() or "nucleotide sequence"


def fetch_ncbi_identification(sequence_id: str, sequence: str):
    result = {
        "query": sequence_id,
        "accession": sequence_id if sequence_id != "raw-sequence" else None,
        "gene": None,
        "organism": None,
        "length": len(sequence),
        "e_value": None,
        "identity_percent": None,
        "tax_id": None,
    }
    try:
        Entrez.email = "developer@example.com"
        with Entrez.esearch(db="nuccore", term=sequence_id, retmax=1) as search_handle:
            search_result = Entrez.read(search_handle)
        identifiers = search_result.get("IdList", [])
        if not identifiers:
            return {"status": "NCBI identification data unavailable", **result}
        with Entrez.efetch(db="nuccore", id=identifiers[0], rettype="gb", retmode="text") as fetch_handle:
            record = SeqIO.read(fetch_handle, "genbank")
        result["accession"] = record.id or result["accession"]
        result["length"] = len(record.seq)
        result["organism"] = record.annotations.get("organism")
        for dbxref in record.dbxrefs:
            if dbxref.startswith("taxon:"):
                result["tax_id"] = dbxref.split(":", 1)[1]
                break
        for feature in record.features:
            if feature.type in {"gene", "CDS"}:
                qualifiers = feature.qualifiers
                result["gene"] = (qualifiers.get("gene") or qualifiers.get("locus_tag") or [None])[0]
                if result["gene"]:
                    break
        return result
    except Exception as exc:
        return {"status": "NCBI identification data unavailable", "error": str(exc), **result}


def fetch_literature_data(query: str):
    records, _ = fetch_online_open_access_context(query)
    return [
        record for record in records
        if record.get("database") == "Europe PMC"
        and (record.get("abstract") or record.get("title"))
    ]


def _uniprot_result(response):
    response.raise_for_status()
    payload = response.json()
    return payload.get("results", []) if "results" in payload else [payload]


def uniprot_fetch_by_accession(accession_id: str):
    response = requests.get(
        UNIPROT_ENTRY_ENDPOINT.format(accession_id=accession_id.strip()),
        timeout=API_TIMEOUT_SECONDS,
    )
    results = _uniprot_result(response)
    return results[0] if results else None


def uniprot_fetch_by_gene(gene_name: str):
    response = requests.get(
        UNIPROT_GENE_ENDPOINT.format(gene_name=requests.utils.quote(gene_name.strip())),
        timeout=API_TIMEOUT_SECONDS,
    )
    results = _uniprot_result(response)
    return results[0] if results else None


def uniprot_fetch_by_organism(tax_id: str):
    response = requests.get(
        UNIPROT_ORGANISM_ENDPOINT.format(tax_id=str(tax_id).strip()),
        timeout=API_TIMEOUT_SECONDS,
    )
    results = _uniprot_result(response)
    return results[0] if results else None


def uniprot_fetcher(query: str, gene_name: str = "", tax_id: str = ""):
    try:
        with st.spinner("Fetching UniProt data..."):
            accession_match = re.fullmatch(
                r"(?:[OPQ][0-9][A-Z0-9]{3}[0-9]|[A-NR-Z][0-9][A-Z0-9]{3}[0-9])",
                query.strip().upper(),
            )
            entry = uniprot_fetch_by_accession(query) if accession_match else None
            if not entry:
                entry = uniprot_fetch_by_gene(gene_name or query)
            organism_entry = uniprot_fetch_by_organism(tax_id) if tax_id else None
            if not entry and not organism_entry:
                return {"status": "UniProt data unavailable", "query": query}
            entry = entry or organism_entry
            domains = [
                feature.get("description", "")
                for feature in entry.get("features", [])
                if feature.get("type") in {"Domain", "Region"}
            ]
            locations = [
                comment.get("texts", [{}])[0].get("value", "")
                for comment in entry.get("comments", [])
                if comment.get("commentType") == "SUBCELLULAR LOCATION"
            ]
            functions = [
                text.get("value", "")
                for comment in entry.get("comments", [])
                if comment.get("commentType") == "FUNCTION"
                for text in comment.get("texts", [])
            ]
            return {
                "accession": entry.get("primaryAccession"),
                "protein": entry.get("proteinDescription", {}).get("recommendedName", {}).get("fullName", {}).get("value"),
                "function": functions,
                "subcellular_location": locations,
                "functional_domains": domains,
            }
    except Exception as exc:
        return {"status": "UniProt data unavailable", "error": str(exc)}


def clinvar_fetcher(query: str):
    try:
        Entrez.email = "developer@example.com"
        with st.spinner("Fetching ClinVar data..."):
            search_handle = Entrez.esearch(db="clinvar", term=query, retmax=5)
            search_result = Entrez.read(search_handle)
            search_handle.close()
            identifiers = search_result.get("IdList", [])
            if not identifiers:
                return {"status": "ClinVar data unavailable", "query": query}
            fetch_handle = Entrez.efetch(db="clinvar", id=identifiers, rettype="vcv", retmode="xml")
            raw_xml = fetch_handle.read()
            fetch_handle.close()
            return {"query": query, "record_ids": identifiers, "summary": str(raw_xml)[:12000]}
    except Exception as exc:
        return {"status": "ClinVar data unavailable", "error": str(exc)}


def kegg_fetcher(query: str):
    try:
        with st.spinner("Fetching KEGG pathway data..."):
            response = requests.get(
                "[https://rest.kegg.jp/find/genes](https://rest.kegg.jp/find/genes)",
                params={"term": query},
                timeout=API_TIMEOUT_SECONDS,
            )
            response.raise_for_status()
            matches = [line.split("\t", 1) for line in response.text.splitlines() if "\t" in line]
            if not matches:
                return {"status": "KEGG data unavailable", "query": query}
            gene_id = matches[0][0]
            link_response = requests.get(
                f"[https://rest.kegg.jp/link/pathway/](https://rest.kegg.jp/link/pathway/){gene_id}",
                timeout=API_TIMEOUT_SECONDS,
            )
            link_response.raise_for_status()
            pathways = [line.split("\t", 1)[1] for line in link_response.text.splitlines() if "\t" in line]
            return {"query": query, "gene_match": matches[0][1], "pathways": pathways}
    except Exception as exc:
        return {"status": "KEGG data unavailable", "error": str(exc)}


def pdb_fetcher(query: str):
    try:
        with st.spinner("Fetching RCSB PDB data..."):
            response = requests.post(
                "[https://search.rcsb.org/rcsbsearch/v2/query](https://search.rcsb.org/rcsbsearch/v2/query)",
                json={
                    "query": {"type": "terminal", "service": "full_text", "parameters": {"value": query}},
                    "return_type": "entry",
                    "request_options": {"paginate": {"start": 0, "rows": 5}},
                },
                timeout=API_TIMEOUT_SECONDS,
            )
            response.raise_for_status()
            identifiers = [item.get("identifier") for item in response.json().get("result_set", [])]
            if not identifiers:
                return {"status": "RCSB PDB data unavailable", "query": query}
            entries = []
            for pdb_id in identifiers:
                entry_response = requests.get(
                    f"[https://data.rcsb.org/rest/v1/core/entry/](https://data.rcsb.org/rest/v1/core/entry/){pdb_id}",
                    timeout=API_TIMEOUT_SECONDS,
                )
                entry_response.raise_for_status()
                entry = entry_response.json().get("struct", {})
                entries.append({
                    "pdb_id": pdb_id,
                    "title": entry.get("title", ""),
                    "experimental_method": entry.get("pdbx_descriptor", ""),
                })
            return {"query": query, "pdb_ids": identifiers, "structural_summary": entries}
    except Exception as exc:
        return {"status": "RCSB PDB data unavailable", "error": str(exc)}


def generate_txt(content: str) -> str:
    return content


def generate_docx(content: str) -> bytes:
    document = WordDocument()
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 3)
            document.add_heading(stripped.lstrip("# "), level=level)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            document.add_paragraph(stripped)
    from io import BytesIO
    binary_output = BytesIO()
    document.save(binary_output)
    return binary_output.getvalue()


def build_attachment_content(uploaded_files, camera_image):
    content = []
    attachment_names = []
    for uploaded_file in uploaded_files or []:
        attachment_names.append(uploaded_file.name)
        file_bytes = uploaded_file.getvalue()
        mime_type = uploaded_file.type or "application/octet-stream"
        if mime_type.startswith("image/"):
            encoded = base64.b64encode(file_bytes).decode("ascii")
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{mime_type};base64,{encoded}"},
            })
        elif mime_type == "application/pdf" or uploaded_file.name.lower().endswith(".pdf"):
            pages = PdfReader(BytesIO(file_bytes)).pages
            content.append("\n".join(page.extract_text() or "" for page in pages))
        else:
            content.append(file_bytes.decode("utf-8", errors="replace"))
    if camera_image is not None:
        attachment_names.append("camera-capture.jpg")
        encoded = base64.b64encode(camera_image.getvalue()).decode("ascii")
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{camera_image.type or 'image/jpeg'};base64,{encoded}"},
        })
    return content, attachment_names


def analyze_bio_context_with_attachments(metrics, attachments, query):
    existing_report = st.session_state.get("bio_report", "")
    attachment_parts, attachment_names = attachments
    prompt = f"""
คุณคือ Expert Bioinformatics Research Analyst
วิเคราะห์ข้อมูลเดิมและไฟล์แนบที่ผู้ใช้ส่งมา โดยแยกข้อเท็จจริงจากการอนุมาน
อย่าสร้างข้อมูลหรือ citation ที่ไม่มีหลักฐาน และแจ้งข้อจำกัดเมื่อข้อมูลไม่พอ

ผลวิเคราะห์ Bioinformatics เดิม:
{existing_report or '(ยังไม่มีรายงานเดิม ใช้ deterministic metrics เป็นหลัก)'}

Deterministic metrics:
{json.dumps(metrics, ensure_ascii=False, indent=2)}

ชื่อไฟล์แนบ: {', '.join(attachment_names) or '(ไม่มีไฟล์แนบ)'}
คำถามเพิ่มเติม: {query or 'วิเคราะห์ข้อมูลทั้งหมดและสรุปประเด็นสำคัญ'}
"""
    message_content = [prompt, *attachment_parts]
    response = get_active_llm(
        st.session_state.active_model_choice,
        st.session_state.active_model_key,
    ).invoke([HumanMessage(content=message_content)])
    return extract_text(response)


def render_bioinformatics():
    st.header("🧬 Bioinformatics Sequence Agent")
    st.caption("Biopython metrics first, Gemini interpretation second")
    input_method = st.selectbox(
        "แหล่ง sequence",
        ["File Upload", "Raw Text Input", "NCBI Accession"],
        key="sequence_input_method",
    )
    uploaded_file = None
    raw_input = ""
    accession = ""
    if input_method == "File Upload":
        uploaded_file = st.file_uploader("อัปโหลด FASTA", type=["fasta", "fa"], key="sequence_file")
    elif input_method == "Raw Text Input":
        raw_input = st.text_area("Nucleotide sequence", height=140, key="sequence_text")
    else:
        accession = st.text_input("NCBI Accession Number", placeholder="เช่น NM_000546 หรือ NC_045512", key="ncbi_accession")
    query = st.text_input("คำถามสำหรับนักชีวสารสนเทศ", placeholder="วิเคราะห์หน้าที่หรือ biological implications ของ sequence นี้")
    analysis_submitted = st.button("วิเคราะห์โครงสร้าง DNA และสร้างรายงาน", type="primary")
    if analysis_submitted:
        try:
            if input_method == "File Upload":
                if uploaded_file is None:
                    raise ValueError("กรุณาเลือกไฟล์ FASTA")
                sequence, sequence_id = parse_uploaded_sequence(uploaded_file)
            elif input_method == "Raw Text Input":
                sequence, sequence_id = parse_raw_sequence(raw_input)
            else:
                sequence, sequence_id = fetch_ncbi_sequence(accession)
            st.session_state.bio_metrics = calculate_sequence_metrics(sequence)
            st.session_state.bio_sequence_id = sequence_id
        except ValueError as exc:
            st.error(str(exc))
            return
    metrics = st.session_state.get("bio_metrics")
    if not metrics:
        return
    st.subheader(f"ผล deterministic analysis: {st.session_state.get('bio_sequence_id', '')}")
    col1, col2, col3 = st.columns(3)
    col1.metric("Sequence length", f"{metrics['length']} bp")
    col2.metric("GC-content", f"{metrics['gc']:.2f}%")
    col3.metric("Protein length", f"{len(metrics['protein'])} aa")
    with st.expander("ดูผลการถอดรหัส"):
        st.code(f"mRNA:\n{metrics['mrna']}\n\nProtein (to_stop=True):\n{metrics['protein'] or '(ไม่พบโปรตีนก่อน stop codon)'}")
    if not GOOGLE_API_KEY:
        st.warning("ตั้งค่า GOOGLE_API_KEY เพื่อใช้การวิเคราะห์ด้วย Gemini")
        return
    if analysis_submitted:
        database_query = resolve_database_query(
            st.session_state.get("bio_sequence_id", ""),
            query,
            metrics,
        )
        identification_data = fetch_ncbi_identification(
            database_query,
            metrics["sequence"],
        )
        protein_data = {
            "UniProt": uniprot_fetcher(
                database_query,
                gene_name=identification_data.get("gene") or "",
                tax_id=identification_data.get("tax_id") or "",
            ),
            "ClinVar": clinvar_fetcher(database_query),
            "KEGG": kegg_fetcher(database_query),
            "RCSB PDB": pdb_fetcher(database_query),
        }
        literature_data = fetch_literature_data(database_query)
        prompt = f"""นี่คือข้อมูลที่ระบบดึงมาจากฐานข้อมูล Open Access:
    ข้อกำหนด: ใช้ข้อมูลจากสาม block ด้านล่างเท่านั้น ห้ามใช้ pretrained knowledge หรือเดา
    ทุก factual claim และ inference ต้องมี citation identifier จาก context ต่อท้ายทันที
    แยกหัวข้อ Facts (ข้อเท็จจริง) และ Inferences (การอนุมาน) อย่างชัดเจน
    หาก block ใดว่าง ให้ระบุว่า ไม่พบข้อมูลในฐานข้อมูล [ชื่อฐานข้อมูล]
    ตอบเป็น Academic Thai และห้ามสร้าง citation หรือ identifier ขึ้นเอง
    <open_access_context>
<identification_data>
{json.dumps(identification_data, ensure_ascii=False, indent=2)}
</identification_data>
<protein_data>
{json.dumps(protein_data, ensure_ascii=False, indent=2)}
</protein_data>
<literature_data>
{json.dumps(literature_data, ensure_ascii=False, indent=2)}
</literature_data>
</open_access_context>

จงวิเคราะห์ข้อมูลทั้งหมดนี้ตามคำสั่ง System Prompt อย่างเคร่งครัด
User Query: {query or 'วิเคราะห์หน้าที่และความสำคัญทางชีวภาพของ sequence นี้'}
"""
        try:
            with st.spinner("กำลังสังเคราะห์ข้อมูลจากฐานข้อมูลภายนอกด้วย Gemini..."):
                response = get_active_llm(
                    st.session_state.active_model_choice,
                    st.session_state.active_model_key,
                ).invoke([HumanMessage(content=prompt)])
            report = extract_text(response)
            st.session_state.bio_report = report
        except Exception:
            st.error("Gemini วิเคราะห์ sequence ไม่สำเร็จ กรุณาตรวจสอบ API key หรือ quota")

    report = st.session_state.get("bio_report")
    if report:
        st.divider()
        st.subheader("Academic Bioinformatics Report")
        st.markdown(report)
        st.download_button(
            "💾 ดาวน์โหลดรายงาน (Word/DOCX)",
            data=generate_docx(report),
            file_name="bioinformatics_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        st.download_button(
            "📄 ดาวน์โหลดรายงาน (TXT)",
            data=generate_txt(report),
            file_name="bioinformatics_report.txt",
            mime="text/plain",
        )


# =============================================================================
# State router: each mode owns its own session-state keys
# =============================================================================

MODE_OPTIONS = ["Online Open Access Research", "Bioinformatics Agent"]
if "active_mode" not in st.session_state or st.session_state.active_mode == "Local Document RAG":
    st.session_state.active_mode = MODE_OPTIONS[0]
with st.sidebar:
    st.title("AI Research Workbench")
    st.session_state.active_model_choice = st.selectbox(
        "เลือกโมเดล AI (Free Tier)",
        list(MODEL_OPTIONS),
        key="model_selector",
    )
    active_model_config = MODEL_OPTIONS[st.session_state.active_model_choice]
    configured_key = get_configured_key(active_model_config["secret"])
    if configured_key:
        st.session_state.active_model_key = configured_key
    else:
        api_key_state = f"api_key_{active_model_config['secret']}"
        st.session_state.active_model_key = st.text_input(
            "Enter API Key",
            type="password",
            key=api_key_state,
        ).strip()
        if not st.session_state.active_model_key:
            st.warning(f"ยังไม่มี {active_model_config['secret']} สำหรับโมเดลที่เลือก")
    st.session_state.active_mode = st.radio("โหมดการทำงาน", MODE_OPTIONS, index=MODE_OPTIONS.index(st.session_state.active_mode))
    st.divider()
    st.caption(f"Active model: {active_model_config['model']}")

if st.session_state.active_mode == "Online Open Access Research":
    render_document_rag()
else:
    render_bioinformatics()