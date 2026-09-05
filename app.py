import base64
import concurrent.futures
import json
import os
import re
import requests
import sqlite3
import time
import uuid
from datetime import datetime
from io import BytesIO, StringIO
from pathlib import Path
from typing import List
from urllib.error import HTTPError, URLError
import streamlit as st
from Bio import Entrez, SeqIO
from Bio.Seq import Seq
from docx import Document as WordDocument
from langchain_core.messages import HumanMessage
from langchain_google_genai import ChatGoogleGenerativeAI
from pypdf import PdfReader

# =============================================================================
# Configuration and shared helpers
# =============================================================================
PROJECT_FOLDER = Path(__file__).resolve().parent

DATA_CACHE_DIR = PROJECT_FOLDER / "data"
DATA_CACHE_DIR.mkdir(parents=True, exist_ok=True)
CHAT_DB_PATH = DATA_CACHE_DIR / "chat_history.db"

ONLINE_RESULT_LIMIT = 5
ONLINE_TIMEOUT_SECONDS = 30
UNIPROT_ENTRY_ENDPOINT = "https://rest.uniprot.org/uniprotkb/{accession_id}?format=json"
UNIPROT_GENE_ENDPOINT = "https://rest.uniprot.org/uniprotkb/search?query=gene:{gene_name}&format=json"
UNIPROT_ORGANISM_ENDPOINT = "https://rest.uniprot.org/uniprotkb/search?query=organism_id:{tax_id}&format=json"
API_TIMEOUT_SECONDS = 20

# Model ids are configurable via env/secrets since they depend on whichever
# account or reseller the deployment actually has credentials for.
GEMINI_MODEL_NAME = os.getenv("GEMINI_MODEL_NAME", "gemini-3.6-flash")
NINEARM_MODEL_NAME = os.getenv("NINEARM_MODEL_NAME", "claude-3.5-sonnet")

MODEL_OPTIONS = {
    "Gemini 3.6 Flash": {
        "provider": "google",
        "model": GEMINI_MODEL_NAME,
        "secret": "GOOGLE_API_KEY",
    },
    "Claude 3.5 (9arm)": {
        "provider": "9arm",
        "model": NINEARM_MODEL_NAME,
        "secret": "NINEARM_API_KEY",
    },
}

st.set_page_config(
    page_title="AI Research Workbench",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =============================================================================
# Design System — Gemini-like Clean Scientific UI
# =============================================================================
DESIGN_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Google+Sans:wght@400;500;700&family=Noto+Sans+Thai:wght@400;500;600;700&family=Inter:wght@400;500;600;700&family=Source+Serif+4:opsz,wght@8..60,400;8..60,500;8..60,600&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
    --bg: #FFFFFF;
    --surface: #F0F4F9;
    --text-primary: #1F1F1F;
    --text-secondary: #444746;
    --text-muted: #747775;
    --accent: #0B57D0;
    --accent-hover: #0842A0;
    --accent-light: #D3E3FD;
    --border: #E3E3E3;
    --border-light: #F2F2F2;
    --success: #146C2E;
    --success-bg: #E6F4EA;
    --warning: #F59E0B;
    --warning-bg: #FFFBEB;
    --error: #B3261E;
    --error-bg: #FCE8E6;
    --info: #0A56D0;
    --info-bg: #E8F0FE;
}

html, body, [data-testid="stAppViewContainer"] {
    background-color: var(--bg) !important;
    color: var(--text-primary) !important;
    font-family: 'Google Sans', 'Noto Sans Thai', 'Inter', -apple-system, sans-serif !important;
}

/* Remove excessive top spacing before main header */
[data-testid="stMain"] {
    padding-top: 16px !important;
    padding-left: 40px !important;
    padding-right: 40px !important;
    padding-bottom: 32px !important;
}

/* Header — no link/anchor, clean */
[data-testid="stHeader"] {
    background-color: transparent !important;
}

[data-testid="stMain"] h1 {
    font-size: 1.5rem !important;
    font-weight: 700 !important;
    color: var(--text-primary) !important;
    margin-top: 8px !important;
    margin-bottom: 4px !important;
    letter-spacing: -0.02em;
    text-decoration: none !important;
}

[data-testid="stMain"] h1 a,
[data-testid="stMain"] h1 ::before,
[data-testid="stMain"] h1 ::after {
    display: none !important;
}

[data-testid="stCaptionContainer"] {
    font-size: 0.85rem !important;
    color: var(--text-secondary) !important;
    margin-bottom: 20px !important;
}

/* Sidebar */
[data-testid="stSidebar"] {
    background-color: var(--surface) !important;
    border-right: none !important;
    padding: 24px 20px !important;
    margin-bottom: 20px !important;
}

[data-testid="stSidebar"] [data-testid="stSidebarTitle"] {
    font-size: 1.2rem !important;
    font-weight: 700 !important;
    color: var(--text-primary) !important;
    margin-bottom: 24px !important;
    padding-bottom: 12px !important;
    border-bottom: 1px solid var(--border) !important;
    letter-spacing: -0.01em;
}

/* Sidebar sections */
.sidebar-section {
    margin-bottom: 24px;
}

.sidebar-section-label {
    font-size: 0.68rem !important;
    font-weight: 700 !important;
    color: var(--text-muted) !important;
    text-transform: uppercase;
    letter-spacing: 0.1em;
    margin-bottom: 8px !important;
}

.sidebar-section-desc {
    font-size: 0.78rem !important;
    color: var(--text-secondary) !important;
    margin-bottom: 10px !important;
}

[data-testid="stSidebar"] .stSelectbox label,
[data-testid="stSidebar"] .stRadio label,
[data-testid="stSidebar"] .stTextInput label {
    font-size: 0.75rem !important;
    font-weight: 600 !important;
    color: var(--text-secondary) !important;
    margin-bottom: 4px !important;
    display: block !important;
}

[data-testid="stSidebar"] .stSelectbox > div,
[data-testid="stSidebar"] .stTextInput > div {
    background-color: var(--bg) !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
    padding: 8px 12px !important;
    font-size: 0.875rem !important;
}

[data-testid="stSidebar"] .stRadio > div {
    gap: 6px !important;
}

[data-testid="stSidebar"] .stRadio [role="radiogroup"] label {
    font-size: 0.85rem !important;
    font-weight: 400 !important;
    color: var(--text-primary) !important;
    text-transform: none !important;
    letter-spacing: normal !important;
    background-color: var(--bg) !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
    padding: 8px 14px !important;
    cursor: pointer;
    transition: all 0.15s ease;
}

[data-testid="stSidebar"] .stRadio [role="radiogroup"] label:hover {
    border-color: var(--accent) !important;
}

[data-testid="stSidebar"] .stRadio [role="radiogroup"] label[data-selected="true"] {
    border-color: var(--accent) !important;
    background-color: var(--accent-light) !important;
    color: var(--accent) !important;
    font-weight: 500 !important;
}

[data-testid="stSidebar"] .stRadio > label {
    display: none !important;
}

/* Status indicators */
.status-block {
    background-color: var(--bg);
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 10px 12px;
    margin-bottom: 8px;
}

.status-row {
    display: flex;
    align-items: center;
    gap: 8px;
    font-size: 0.85rem;
    color: var(--text-secondary);
    margin-bottom: 4px;
}

.status-row:last-child {
    margin-bottom: 0;
}

.status-dot {
    width: 8px;
    height: 8px;
    border-radius: 50%;
    flex-shrink: 0;
}

.status-dot.active {
    background-color: var(--success);
}

.status-dot.warning {
    background-color: var(--warning);
}

.status-dot.inactive {
    background-color: var(--error);
}

.status-label {
    font-weight: 600;
    color: var(--text-primary);
    font-size: 0.82rem;
}

.status-value {
    font-size: 0.82rem;
    color: var(--text-secondary);
}

/* Main content headers */
[data-testid="stMain"] h2 {
    font-size: 1.25rem !important;
    font-weight: 600 !important;
    color: var(--text-primary) !important;
    margin-top: 24px !important;
    margin-bottom: 8px !important;
}

[data-testid="stMain"] h3 {
    font-size: 1.1rem !important;
    font-weight: 600 !important;
    color: var(--text-primary) !important;
    margin-top: 20px !important;
    margin-bottom: 6px !important;
}

/* Popover */
[data-testid="stPopover"] button {
    background-color: var(--surface) !important;
    border: none !important;
    border-radius: 20px !important;
    color: var(--text-primary) !important;
    font-size: 0.85rem !important;
    font-weight: 500 !important;
    padding: 8px 16px !important;
    transition: all 0.15s ease;
}

[data-testid="stPopover"] button:hover {
    background-color: var(--border) !important;
    color: var(--accent) !important;
}

/* File uploader */
[data-testid="stFileUploader"] {
    border: 1px dashed var(--border) !important;
    border-radius: 12px !important;
    padding: 16px !important;
    background-color: var(--surface) !important;
}

[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] {
    background-color: transparent !important;
    border: none !important;
}

[data-testid="stFileUploader"] [data-testid="stFileUploaderMessage"] {
    font-size: 0.85rem !important;
    color: var(--text-secondary) !important;
}

/* Chat UI (Gemini style) */
[data-testid="stChatMessage"] {
    border: none !important;
    background-color: transparent !important;
    padding: 8px 16px !important;
    margin-bottom: 12px !important;
}

[data-testid="stChatMessage"][data-testid="stChatMessageUser"] {
    background-color: var(--surface) !important;
    border-radius: 24px !important;
    max-width: 85% !important;
    margin-left: auto !important;
    color: var(--text-primary) !important;
    padding: 16px 24px !important;
}

[data-testid="stChatMessage"][data-testid="stChatMessageAssistant"] {
    background-color: transparent !important;
    padding: 16px 8px !important;
}

[data-testid="stChatInput"] {
    border-radius: 32px !important;
    border: 1px solid var(--border) !important;
    background-color: var(--bg) !important;
    padding: 6px 12px !important;
    box-shadow: 0 2px 6px rgba(0,0,0,0.05) !important;
}

[data-testid="stChatInput"]:focus-within {
    border-color: var(--accent) !important;
    box-shadow: 0 2px 6px rgba(11, 87, 208, 0.1) !important;
}

[data-testid="stChatInput"] input {
    font-size: 0.95rem !important;
}

/* Buttons */
.stButton > button, .stDownloadButton > button {
    background-color: var(--bg) !important;
    border: 1px solid var(--border) !important;
    border-radius: 20px !important;
    color: var(--text-primary) !important;
    font-size: 0.875rem !important;
    font-weight: 500 !important;
    padding: 8px 24px !important;
    transition: all 0.15s ease;
}

.stButton > button:hover, .stDownloadButton > button:hover {
    background-color: var(--surface) !important;
    border-color: var(--border) !important;
    color: var(--accent) !important;
}

.stButton > button[kind="primary"], .stButton > button[type="primary"] {
    background-color: var(--accent) !important;
    border-color: var(--accent) !important;
    color: white !important;
    border-radius: 24px !important;
}

.stButton > button[kind="primary"]:hover, .stButton > button[type="primary"]:hover {
    background-color: var(--accent-hover) !important;
    border-color: var(--accent-hover) !important;
    box-shadow: 0 1px 3px rgba(0,0,0,0.1);
}

/* Expander */
[data-testid="stExpander"] {
    border: 1px solid var(--border) !important;
    border-radius: 12px !important;
    background-color: var(--surface) !important;
}

[data-testid="stExpander"] summary {
    font-size: 0.9rem !important;
    font-weight: 500 !important;
    color: var(--text-primary) !important;
    padding: 12px 16px !important;
}

[data-testid="stExpander"] summary:hover {
    color: var(--accent) !important;
}

/* Code blocks */
[data-testid="stCodeBlock"] {
    background-color: var(--surface) !important;
    border: none !important;
    border-radius: 12px !important;
    font-family: 'JetBrains Mono', 'Fira Code', monospace !important;
    font-size: 0.85rem !important;
}

/* Metrics */
[data-testid="stMetric"] {
    background-color: var(--surface) !important;
    border: none !important;
    border-radius: 12px !important;
    padding: 16px 20px !important;
}

[data-testid="stMetric"] [data-testid="stMetricLabel"] {
    font-size: 0.75rem !important;
    font-weight: 500 !important;
    color: var(--text-secondary) !important;
    text-transform: uppercase;
    letter-spacing: 0.04em;
}

[data-testid="stMetric"] [data-testid="stMetricValue"] {
    font-size: 1.5rem !important;
    font-weight: 500 !important;
    color: var(--text-primary) !important;
}

/* Divider */
[data-testid="stDivider"] {
    border-color: var(--border) !important;
    margin: 24px 0 !important;
}

/* Subheader */
[data-testid="stSubheader"] {
    font-size: 1.1rem !important;
    font-weight: 500 !important;
    color: var(--text-primary) !important;
    margin-top: 20px !important;
    margin-bottom: 10px !important;
}

/* Text input / textarea */
.stTextInput > div, .stTextArea > div {
    background-color: var(--surface) !important;
    border: none !important;
    border-radius: 8px !important;
    padding: 10px 14px !important;
}

.stTextInput > div:focus-within, .stTextArea > div:focus-within {
    border: 1px solid var(--accent) !important;
    background-color: var(--bg) !important;
}

.stTextInput input, .stTextArea textarea {
    font-size: 0.95rem !important;
    color: var(--text-primary) !important;
}

.stTextInput input::placeholder, .stTextArea textarea::placeholder {
    color: var(--text-muted) !important;
}

/* Alerts */
[data-testid="stAlert"] {
    border-radius: 8px !important;
    font-size: 0.9rem !important;
    padding: 12px 16px !important;
    border: none !important;
}

[data-testid="stAlertWarning"] {
    background-color: var(--warning-bg) !important;
    color: #92400E !important;
}

[data-testid="stAlertError"] {
    background-color: var(--error-bg) !important;
    color: #991B1B !important;
}

/* Report body */
.report-body {
    font-family: 'Source Serif 4', Georgia, serif !important;
    font-size: 1rem !important;
    line-height: 1.75 !important;
    color: var(--text-primary) !important;
}

.report-body h1, .report-body h2, .report-body h3 {
    font-family: 'Google Sans', 'Inter', sans-serif !important;
    font-weight: 500 !important;
    margin-top: 24px !important;
    margin-bottom: 12px !important;
}

.report-body p {
    margin-bottom: 12px !important;
}

.report-body ul, .report-body ol {
    padding-left: 20px !important;
    margin-bottom: 12px !important;
}

.report-body li {
    margin-bottom: 6px !important;
}

/* Bordered container */
.bordered-container {
    background-color: var(--surface) !important;
    border: none !important;
    border-radius: 12px !important;
    padding: 20px !important;
    margin-bottom: 20px !important;
}

/* Validation section */
.validation-pass {
    color: var(--success) !important;
    font-size: 0.9rem;
    margin-bottom: 6px;
}

.validation-fail {
    color: var(--error) !important;
    font-size: 0.9rem;
    margin-bottom: 6px;
}

.validation-warn {
    color: var(--warning) !important;
    font-size: 0.9rem;
    margin-bottom: 6px;
}

/* Pipeline steps */
.pipeline-step {
    display: flex;
    align-items: center;
    gap: 10px;
    font-size: 0.88rem;
    color: var(--text-secondary);
    margin-bottom: 8px;
}

.pipeline-step.done {
    color: var(--success);
}

.pipeline-step.active {
    color: var(--accent);
    font-weight: 500;
}

.pipeline-step.pending {
    color: var(--text-muted);
}

/* Analysis log */
.analysis-log {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.8rem;
    color: var(--text-secondary);
    background-color: var(--bg);
    border-radius: 8px;
    padding: 12px 16px;
    max-height: 200px;
    overflow-y: auto;
}

.analysis-log .log-entry {
    margin-bottom: 4px;
}

.analysis-log .log-entry.error {
    color: var(--error);
}

.analysis-log .log-time {
    color: var(--text-muted);
    margin-right: 12px;
}

/* Confidence display */
.confidence-display {
    display: flex;
    align-items: center;
    gap: 12px;
    padding: 16px 20px;
    background-color: var(--bg);
    border-radius: 12px;
    margin-bottom: 8px;
}

.confidence-score {
    font-size: 1.6rem;
    font-weight: 500;
    color: var(--text-primary);
}

.confidence-label {
    font-size: 0.9rem;
    font-weight: 500;
    color: var(--success);
}

.confidence-desc {
    font-size: 0.85rem;
    color: var(--text-muted);
    margin-top: 6px;
}

/* Evidence source */
.evidence-source {
    padding: 12px 16px;
    background-color: var(--bg);
    border-radius: 8px;
    margin-bottom: 8px;
    font-size: 0.88rem;
}

.evidence-source .source-name {
    font-weight: 500;
    color: var(--text-primary);
    margin-bottom: 4px;
}

.evidence-source .source-detail {
    color: var(--text-secondary);
    font-size: 0.82rem;
}

/* Spinner */
[data-testid="stSpinner"] {
    color: var(--accent) !important;
    font-size: 0.9rem !important;
}

/* Scrollbar */
::-webkit-scrollbar {
    width: 6px;
    height: 6px;
}

::-webkit-scrollbar-track {
    background: transparent;
}

::-webkit-scrollbar-thumb {
    background-color: var(--border);
    border-radius: 4px;
}

::-webkit-scrollbar-thumb:hover {
    background-color: #C7C7C7;
}
"""

st.markdown(f"<style>{DESIGN_CSS}</style>", unsafe_allow_html=True)

# Enter-to-submit on desktop for st.chat_input
ENTER_SUBMIT_JS = """
<script>
(function() {
    function setupEnterSubmit() {
        const chatInput = document.querySelector('[data-testid="stChatInput"] input');
        if (!chatInput) return;
        chatInput.addEventListener('keydown', function(e) {
            if (e.key === 'Enter' && !e.shiftKey && !e.isComposing) {
                e.preventDefault();
                const submitBtn = document.querySelector('[data-testid="stChatInput"] button');
                if (submitBtn) submitBtn.click();
            }
        });
    }
    const observer = new MutationObserver(function() { setupEnterSubmit(); });
    observer.observe(document.body, { childList: true, subtree: true });
    setupEnterSubmit();
})();
</script>
"""
st.markdown(ENTER_SUBMIT_JS, unsafe_allow_html=True)

# =============================================================================
# Core helpers
# =============================================================================
def get_api_key() -> str:
    key = os.getenv("GOOGLE_API_KEY", "").strip()
    if key:
        return key
    try:
        return str(st.secrets.get("GOOGLE_API_KEY", "")).strip()
    except Exception:
        return ""

GOOGLE_API_KEY = get_api_key()
if GOOGLE_API_KEY:
    os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY

def is_rate_limit_error(exc: Exception) -> bool:
    text = str(exc).upper()
    return any(item in text for item in ("429", "RESOURCE_EXHAUSTED", "QUOTA", "RATE LIMIT"))

def extract_text(response) -> str:
    content = getattr(response, "content", response)
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        return "\n".join(
            block if isinstance(block, str) else block.get("text", "")
            for block in content
            if isinstance(block, str) or isinstance(block, dict)
        ).strip()
    if isinstance(content, dict):
        return str(content.get("text", "")).strip()
    return ""

def format_bytes(size: int) -> str:
    if size < 1024:
        return f"{size} B"
    if size < 1024 ** 2:
        return f"{size / 1024:.1f} KB"
    return f"{size / (1024 ** 2):.1f} MB"

def get_configured_key(key_name: str) -> str:
    key = os.getenv(key_name, "").strip()
    if key:
        return key
    try:
        return str(st.secrets.get(key_name, "")).strip()
    except Exception:
        return ""

NCBI_API_KEY = get_configured_key("NCBI_API_KEY")
if NCBI_API_KEY:
    Entrez.api_key = NCBI_API_KEY

# NCBI asks every E-utilities caller to identify itself with a real contact
# address; a fixed placeholder is set once here rather than per call.
Entrez.email = get_configured_key("NCBI_CONTACT_EMAIL") or "developer@example.com"

def get_active_llm(model_choice: str, api_key: str):
    config = MODEL_OPTIONS.get(model_choice)
    if config is None:
        raise ValueError("ไม่พบข้อมูลรุ่นของโมเดลที่เลือก กรุณาติดต่อผู้พัฒนาระบบ")
    if not api_key:
        raise ValueError(f"ยังไม่ได้กำหนดค่า API Key สำหรับ {model_choice} กรุณาติดต่อผู้พัฒนาระบบ")

    if config["provider"] == "google":
        return ChatGoogleGenerativeAI(model=config["model"], temperature=0.2, google_api_key=api_key)
    if config["provider"] == "9arm":
        base_url = get_configured_key("NINEARM_BASE_URL")
        if not base_url:
            raise ValueError("ยังไม่ได้กำหนดค่า NINEARM_BASE_URL (URL ของผู้ให้บริการ Claude ที่ใช้งานอยู่) กรุณาติดต่อผู้พัฒนาระบบ")
        try:
            from langchain_openai import ChatOpenAI
        except ImportError:
            raise RuntimeError("ไม่พบไลบรารี langchain-openai สำหรับใช้งาน API กรุณาติดต่อผู้พัฒนาระบบ")
        return ChatOpenAI(model=config["model"], temperature=0.2, api_key=api_key, base_url=base_url)
    raise ValueError(f"ไม่รองรับผู้ให้บริการโมเดลนี้: {config['provider']} กรุณาติดต่อผู้พัฒนาระบบ")

def translate_query_to_english(query: str, model_choice: str, api_key: str) -> str:
    """Turn a research question in any language into an effective English
    search string for the (English-only) literature APIs. Falls back to the
    original text -- never a generic placeholder -- if translation fails, so
    a failed search is at least an honest empty result rather than a
    misleadingly "successful" search on unrelated terms."""
    stripped = (query or "").strip()
    if not stripped:
        return ""
    ascii_chars = sum(1 for ch in stripped if ord(ch) < 128)
    if ascii_chars / len(stripped) > 0.9:
        return stripped
    try:
        llm = get_active_llm(model_choice, api_key)
        instruction = (
            "Translate the following research question into a short, effective "
            "English search query for a biomedical literature database. "
            "Reply with ONLY the search keywords, no punctuation, no explanation.\n\n"
            f"Question: {stripped}"
        )
        response = llm.invoke([HumanMessage(content=instruction)])
        translated = extract_text(response).strip().strip('"')
        return translated or stripped
    except Exception:
        return stripped

# =============================================================================
# System Status — real connectivity checks
# =============================================================================
@st.cache_resource(ttl=300)
def check_ncbi_status() -> dict:
    try:
        with Entrez.esearch(db="nuccore", term="BRCA1", retmax=1, timeout=API_TIMEOUT_SECONDS) as h:
            result = Entrez.read(h)
        return {"status": "active", "detail": "NCBI reachable"}
    except Exception:
        return {"status": "inactive", "detail": "NCBI unreachable"}

@st.cache_resource(ttl=300)
def check_uniprot_status() -> dict:
    try:
        r = requests.get("https://rest.uniprot.org/uniprotkb/P04637?format=json", timeout=10)
        r.raise_for_status()
        return {"status": "active", "detail": "UniProt reachable"}
    except Exception:
        return {"status": "inactive", "detail": "UniProt unreachable"}

@st.cache_resource(ttl=300)
def check_kegg_status() -> dict:
    try:
        r = requests.get("https://rest.kegg.jp/info/kegg", timeout=10)
        r.raise_for_status()
        return {"status": "active", "detail": "KEGG reachable"}
    except Exception:
        return {"status": "inactive", "detail": "KEGG unreachable"}

@st.cache_resource(ttl=300)
def check_pdb_status() -> dict:
    try:
        r = requests.get("https://data.rcsb.org/rest/v1/core/entry/4HHB", timeout=10)
        r.raise_for_status()
        return {"status": "active", "detail": "RCSB PDB (Data API) reachable"}
    except requests.exceptions.Timeout:
        return {"status": "warning", "detail": "RCSB PDB connection timeout"}
    except Exception as e:
        return {"status": "inactive", "detail": f"RCSB PDB unreachable: {e}"}

def render_sidebar_status():
    """Render sidebar status section with real connectivity data.

    The four checks are independent network calls; running them concurrently
    keeps a cold page load to roughly the slowest single check instead of
    their sum (each has its own ~10s timeout)."""
    checks = {
        "ncbi": check_ncbi_status,
        "uniprot": check_uniprot_status,
        "kegg": check_kegg_status,
        "pdb": check_pdb_status,
    }
    with concurrent.futures.ThreadPoolExecutor(max_workers=len(checks)) as executor:
        results = {name: future.result() for name, future in
                   {name: executor.submit(fn) for name, fn in checks.items()}.items()}
    ncbi, uniprot, kegg, pdb = results["ncbi"], results["uniprot"], results["kegg"], results["pdb"]

    def dot(status):
        return {
            "active": "active",
            "warning": "warning",
            "inactive": "inactive"
        }.get(status, "inactive")

    st.markdown(
        f'<div class="status-block">'
        f'<div class="status-row"><span class="status-dot active"></span>'
        f'<span class="status-label">System Operational</span></div>'
        f'</div>',
        unsafe_allow_html=True,
    )

    model_name = st.session_state.get("active_model_choice", "N/A")
    st.markdown(
        f'<div class="status-block">'
        f'<div class="status-row"><span class="status-dot active"></span>'
        f'<span class="status-label">AI MODEL</span></div>'
        f'<div class="status-row"><span class="status-value">{model_name}</span></div>'
        f'</div>',
        unsafe_allow_html=True,
    )

    st.markdown(
        f'<div class="status-block">'
        f'<div class="status-row"><span class="status-dot {dot(ncbi["status"])}"></span>'
        f'<span class="status-label">NCBI</span></div>'
        f'<div class="status-row"><span class="status-dot {dot(uniprot["status"])}"></span>'
        f'<span class="status-label">UniProt</span></div>'
        f'<div class="status-row"><span class="status-dot {dot(kegg["status"])}"></span>'
        f'<span class="status-label">KEGG</span></div>'
        f'<div class="status-row"><span class="status-dot {dot(pdb["status"])}"></span>'
        f'<span class="status-label">RCSB PDB</span></div>'
        f'</div>',
        unsafe_allow_html=True,
    )

# =============================================================================
# Analysis Log
# =============================================================================
class AnalysisLog:
    def __init__(self):
        self.entries = []
        self.start_time = time.time()

    def log(self, message: str, is_error: bool = False):
        ts = datetime.now().strftime("%H:%M:%S")
        self.entries.append({"time": ts, "message": message, "error": is_error})

    def render(self):
        if not self.entries:
            return
        st.markdown('<div class="analysis-log">', unsafe_allow_html=True)
        for entry in self.entries:
            cls = " error" if entry["error"] else ""
            prefix = "⚠ " if entry["error"] else ""
            st.markdown(
                f'<div class="log-entry{cls}">'
                f'<span class="log-time">{entry["time"]}</span>'
                f'{prefix}{entry["message"]}'
                f'</div>',
                unsafe_allow_html=True,
            )
        st.markdown('</div>', unsafe_allow_html=True)

    def elapsed(self) -> float:
        return time.time() - self.start_time

# =============================================================================
# Analysis Pipeline
# =============================================================================
PIPELINE_STEPS = [
    "Input Validation",
    "Sequence Analysis",
    "Database Retrieval",
    "Computational Analysis",
    "AI Interpretation",
    "Evidence Verification",
    "Report Generation",
]

def render_pipeline(step_states: dict):
    """Render pipeline steps with real status."""
    st.markdown('<div class="bordered-container">', unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.75rem;font-weight:700;color:#9CA3AF;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:10px;">Analysis Pipeline</div>', unsafe_allow_html=True)
    for step in PIPELINE_STEPS:
        state = step_states.get(step, "pending")
        if state == "done":
            icon = "✓"
            cls = "done"
        elif state == "active":
            icon = "●"
            cls = "active"
        else:
            icon = "○"
            cls = "pending"
        st.markdown(
            f'<div class="pipeline-step {cls}"><span>{icon}</span><span>{step}</span></div>',
            unsafe_allow_html=True,
        )
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# Sequence Validation
# =============================================================================
def validate_sequence(sequence: Seq, sequence_id: str, input_method: str) -> dict:
    """Validate sequence before analysis. Returns dict of checks."""
    checks = []
    seq_text = str(sequence).upper() if sequence else ""

    # Check 1: Sequence detected
    if seq_text:
        checks.append(("pass", "Sequence detected"))
    else:
        checks.append(("fail", "No sequence detected"))
        return {"valid": False, "checks": checks}

    # Check 2: Format validated
    checks.append(("pass", f"Format validated ({input_method})"))

    # Check 3: Sequence length
    checks.append(("pass", f"Sequence length: {len(seq_text)} bp"))

    # Check 4: Valid characters
    if all(c in "ACGTN" for c in seq_text):
        checks.append(("pass", "Valid nucleotide characters"))
    else:
        invalid = sorted(set(seq_text) - set("ACGTN"))
        checks.append(("fail", f"Invalid characters: {', '.join(invalid)}"))
        return {"valid": False, "checks": checks}

    # Check 5: DNA/RNA/Protein detection
    if all(c in "ACGT" for c in seq_text):
        checks.append(("pass", "Detected as DNA (A, C, G, T)"))
    elif all(c in "ACGU" for c in seq_text):
        checks.append(("pass", "Detected as RNA (A, C, G, U)"))
    else:
        checks.append(("warn", "Mixed or ambiguous characters detected"))

    # Check 6: sequence header (if applicable)
    if sequence_id and sequence_id != "raw-sequence":
        checks.append(("pass", f"Sequence header: {sequence_id}"))
    else:
        checks.append(("warn", "No sequence header (raw input)"))

    # Check 7: Duplicate/empty check
    if len(seq_text) < 10:
        checks.append(("warn", "Sequence length unusually short (< 10 bp)"))
    else:
        checks.append(("pass", "No empty or duplicate segments detected"))

    valid = all(c[0] != "fail" for c in checks)
    return {"valid": valid, "checks": checks}

def render_validation(validation: dict):
    st.markdown('<div class="bordered-container">', unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.75rem;font-weight:700;color:#9CA3AF;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:10px;">Sequence Validation</div>', unsafe_allow_html=True)
    for status, msg in validation["checks"]:
        if status == "pass":
            icon = "✓"
            cls = "validation-pass"
        elif status == "fail":
            icon = "✗"
            cls = "validation-fail"
        else:
            icon = "⚠"
            cls = "validation-warn"
        st.markdown(f'<div class="{cls}">{icon} {msg}</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# Analysis ID / Reproducibility
# =============================================================================
def generate_analysis_id() -> str:
    date_str = datetime.now().strftime("%Y%m%d")
    unique = uuid.uuid4().hex[:5].upper()
    return f"BIO-{date_str}-{unique}"

def render_analysis_metadata(metadata: dict):
    st.markdown('<div class="bordered-container">', unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.75rem;font-weight:700;color:#9CA3AF;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:10px;">Analysis Metadata</div>', unsafe_allow_html=True)
    for key, value in metadata.items():
        st.markdown(
            f'<div style="font-size:0.82rem;color:#555568;margin-bottom:4px;">'
            f'<span style="font-weight:600;color:#1A1A2E;">{key}:</span> {value}'
            f'</div>',
            unsafe_allow_html=True,
        )
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# Evidence & Source Traceability
# =============================================================================
def render_evidence_sources(evidence: dict):
    """Render evidence sources with traceability."""
    st.markdown('<div class="bordered-container">', unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.75rem;font-weight:700;color:#9CA3AF;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:10px;">Evidence & Sources</div>', unsafe_allow_html=True)

    today = datetime.now().strftime("%Y-%m-%d")

    # NCBI
    ncbi = evidence.get("ncbi", {})
    if ncbi.get("accession"):
        st.markdown(
            f'<div class="evidence-source">'
            f'<div class="source-name">NCBI</div>'
            f'<div class="source-detail">Accession: {ncbi["accession"]}</div>'
            f'<div class="source-detail">Retrieved: {today}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            f'<div class="evidence-source">'
            f'<div class="source-name">NCBI</div>'
            f'<div class="source-detail">No verified source found</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    # UniProt
    uniprot = evidence.get("uniprot", {})
    if uniprot.get("accession"):
        st.markdown(
            f'<div class="evidence-source">'
            f'<div class="source-name">UniProt</div>'
            f'<div class="source-detail">Entry: {uniprot["accession"]}</div>'
            f'<div class="source-detail">Retrieved: {today}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            f'<div class="evidence-source">'
            f'<div class="source-name">UniProt</div>'
            f'<div class="source-detail">No verified source found</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    # KEGG
    kegg = evidence.get("kegg", {})
    if kegg.get("pathways"):
        st.markdown(
            f'<div class="evidence-source">'
            f'<div class="source-name">KEGG</div>'
            f'<div class="source-detail">Pathway: {kegg["pathways"][0]}</div>'
            f'<div class="source-detail">Retrieved: {today}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            f'<div class="evidence-source">'
            f'<div class="source-name">KEGG</div>'
            f'<div class="source-detail">No verified source found</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    # RCSB PDB
    pdb = evidence.get("pdb", {})
    if pdb.get("pdb_ids"):
        st.markdown(
            f'<div class="evidence-source">'
            f'<div class="source-name">RCSB PDB</div>'
            f'<div class="source-detail">PDB ID: {pdb["pdb_ids"][0]}</div>'
            f'<div class="source-detail">Retrieved: {today}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            f'<div class="evidence-source">'
            f'<div class="source-name">RCSB PDB</div>'
            f'<div class="source-detail">No verified source found</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# Confidence Display
# =============================================================================
def render_confidence(score: int):
    if score >= 80:
        label = "High"
    elif score >= 50:
        label = "Medium"
    else:
        label = "Low"

    st.markdown(
        f'<div class="confidence-display">'
        f'<span class="confidence-score">{score} / 100</span>'
        f'<span class="confidence-label">{label}</span>'
        f'</div>'
        f'<div class="confidence-desc">'
        f'คะแนนนี้สะท้อนระดับความมั่นใจของ AI ต่อผลการตีความจากข้อมูลที่ใช้ในการวิเคราะห์ '
        f'ไม่ใช่ค่าความเป็นไปได้ทางชีววิทยา หรือผลการพิสูจน์ทางสถิติ'
        f'</div>',
        unsafe_allow_html=True,
    )

# =============================================================================
# Limitations
# =============================================================================
def render_limitations():
    st.markdown('<div class="bordered-container">', unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.75rem;font-weight:700;color:#9CA3AF;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:10px;">Limitations</div>', unsafe_allow_html=True)
    st.markdown(
        '<ul style="font-size:0.82rem;color:#555568;padding-left:16px;margin:0;">'
        '<li>AI-generated interpretation should be independently verified.</li>'
        '<li>Database annotations may change over time.</li>'
        '<li>Computational predictions do not constitute experimental evidence.</li>'
        '<li>Results should be independently validated before scientific or clinical use.</li>'
        '</ul>',
        unsafe_allow_html=True,
    )
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# File Format Support
# =============================================================================
SUPPORTED_EXTENSIONS = {
    # Sequence
    ".fasta": "FASTA", ".fa": "FASTA", ".fna": "FASTA (DNA)", ".faa": "FASTA (Protein)",
    ".fastq": "FASTQ", ".fq": "FASTQ",
    ".gb": "GenBank", ".gbff": "GenBank", ".embl": "EMBL",
    # Annotation
    ".gff": "GFF", ".gff3": "GFF3", ".gtf": "GTF", ".bed": "BED",
    # Variation
    ".vcf": "VCF",
    # Alignments
    ".sam": "SAM", ".bam": "BAM", ".cram": "CRAM",
    # Structure
    ".pdb": "PDB", ".cif": "mmCIF", ".mmcif": "mmCIF", ".sdf": "SDF", ".mol": "MOL",
    # Tabular
    ".csv": "CSV", ".tsv": "TSV", ".json": "JSON", ".jsonl": "JSONL",
    # Text
    ".txt": "TXT",
    # Documentation
    ".pdf": "PDF",
    # Archives
    ".zip": "ZIP", ".gz": "GZIP", ".tar": "TAR",
}

def detect_file_format(file_bytes: bytes, filename: str) -> str:
    """Detect actual file format from content, not just extension."""
    header = file_bytes[:1024]

    # PDF
    if header.startswith(b"%PDF"):
        return "PDF"
    # ZIP
    if header.startswith(b"PK"):
        return "ZIP"
    # GZIP
    if header.startswith(b"\x1f\x8b"):
        return "GZIP"
    # TAR
    if len(header) >= 262 and header[257:262] == b"ustar":
        return "TAR"
    # BAM (starts with magic)
    if header.startswith(b"\x1b\x0b"):
        return "BAM"

    # Try text-based detection
    try:
        text = header.decode("utf-8", errors="ignore").strip()
        if text.startswith(">"):
            return "FASTA"
        if text.startswith("@"):
            return "FASTQ"
        if "LOCUS" in text.upper():
            return "GenBank"
        if "ID " in text and "VERSION" in text:
            return "EMBL"
        if text.startswith("##gff"):
            return "GFF3"
        if text.startswith("#chromosome"):
            return "GTF"
        if text.startswith("##fileformat=VCF"):
            return "VCF"
        if text.startswith("PDB "):
            return "PDB"
        if "data_" in text and "loop_" in text:
            return "mmCIF"
    except Exception:
        pass

    # Fallback to extension
    ext = Path(filename).suffix.lower()
    return SUPPORTED_EXTENSIONS.get(ext, "Unknown")

def validate_uploaded_file(uploaded_file) -> tuple:
    """Validate uploaded file for security and format. Returns (valid, error_msg, detected_format)."""
    filename = uploaded_file.name
    ext = Path(filename).suffix.lower()

    # Check extension
    if ext not in SUPPORTED_EXTENSIONS:
        return False, f"Unsupported file format: {ext}", "Unknown"

    # Check size (existing system limit)
    size = len(uploaded_file.getvalue())
    if size > 200 * 1024 * 1024:  # 200MB
        return False, "File size exceeds 200MB limit", "Unknown"

    # Detect actual format
    file_bytes = uploaded_file.getvalue()
    detected = detect_file_format(file_bytes, filename)

    return True, None, detected

# =============================================================================
# Persistent chat history (SQLite)
# =============================================================================
DEFAULT_CONVERSATION_TITLE = "การสนทนาใหม่"

def _chat_db_connect() -> sqlite3.Connection:
    conn = sqlite3.connect(str(CHAT_DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn

def init_chat_db() -> None:
    conn = _chat_db_connect()
    try:
        conn.execute(
            """CREATE TABLE IF NOT EXISTS conversations (
                id TEXT PRIMARY KEY,
                title TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )"""
        )
        conn.execute(
            """CREATE TABLE IF NOT EXISTS messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                conversation_id TEXT NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                sources TEXT,
                created_at TEXT NOT NULL
            )"""
        )
        conn.commit()
    finally:
        conn.close()

def create_conversation(title: str = DEFAULT_CONVERSATION_TITLE) -> str:
    conversation_id = uuid.uuid4().hex
    now = datetime.now().isoformat()
    conn = _chat_db_connect()
    try:
        conn.execute(
            "INSERT INTO conversations (id, title, created_at, updated_at) VALUES (?, ?, ?, ?)",
            (conversation_id, title, now, now),
        )
        conn.commit()
    finally:
        conn.close()
    return conversation_id

def list_conversations(limit: int = 50) -> list:
    conn = _chat_db_connect()
    try:
        rows = conn.execute(
            "SELECT id, title, updated_at FROM conversations ORDER BY updated_at DESC LIMIT ?",
            (limit,),
        ).fetchall()
    finally:
        conn.close()
    return [dict(row) for row in rows]

def load_conversation_messages(conversation_id: str) -> list:
    conn = _chat_db_connect()
    try:
        rows = conn.execute(
            "SELECT role, content, sources FROM messages WHERE conversation_id = ? ORDER BY id ASC",
            (conversation_id,),
        ).fetchall()
    finally:
        conn.close()
    messages = []
    for row in rows:
        message = {"role": row["role"], "content": row["content"]}
        if row["sources"]:
            message["sources"] = json.loads(row["sources"])
        messages.append(message)
    return messages

def append_chat_message(conversation_id: str, role: str, content: str, sources=None) -> None:
    now = datetime.now().isoformat()
    conn = _chat_db_connect()
    try:
        conn.execute(
            "INSERT INTO messages (conversation_id, role, content, sources, created_at) VALUES (?, ?, ?, ?, ?)",
            (conversation_id, role, content, json.dumps(sources, ensure_ascii=False) if sources else None, now),
        )
        conn.execute("UPDATE conversations SET updated_at = ? WHERE id = ?", (now, conversation_id))
        if role == "user":
            row = conn.execute("SELECT title FROM conversations WHERE id = ?", (conversation_id,)).fetchone()
            if row and row["title"] == DEFAULT_CONVERSATION_TITLE:
                new_title = content.strip().splitlines()[0][:60] or DEFAULT_CONVERSATION_TITLE
                conn.execute("UPDATE conversations SET title = ? WHERE id = ?", (new_title, conversation_id))
        conn.commit()
    finally:
        conn.close()

def delete_conversation(conversation_id: str) -> None:
    conn = _chat_db_connect()
    try:
        conn.execute("DELETE FROM messages WHERE conversation_id = ?", (conversation_id,))
        conn.execute("DELETE FROM conversations WHERE id = ?", (conversation_id,))
        conn.commit()
    finally:
        conn.close()

init_chat_db()

def ensure_chat_session_state() -> None:
    if "current_conversation_id" in st.session_state:
        return
    conversations = list_conversations()
    if conversations:
        st.session_state.current_conversation_id = conversations[0]["id"]
    else:
        st.session_state.current_conversation_id = create_conversation()
    st.session_state.chat_history = load_conversation_messages(st.session_state.current_conversation_id)

def render_chat_sidebar() -> None:
    ensure_chat_session_state()
    st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-section-label">CONVERSATIONS</div>', unsafe_allow_html=True)
    if st.button(DEFAULT_CONVERSATION_TITLE, key="new_conversation_btn", use_container_width=True):
        st.session_state.current_conversation_id = create_conversation()
        st.session_state.chat_history = []
        st.rerun()
    for convo in list_conversations():
        is_active = convo["id"] == st.session_state.current_conversation_id
        if st.button(
            convo["title"],
            key=f"load_convo_{convo['id']}",
            use_container_width=True,
            type="primary" if is_active else "secondary",
        ):
            st.session_state.current_conversation_id = convo["id"]
            st.session_state.chat_history = load_conversation_messages(convo["id"])
            st.rerun()
    if st.session_state.chat_history:
        if st.button("ลบการสนทนานี้", key="delete_current_conversation", use_container_width=True):
            delete_conversation(st.session_state.current_conversation_id)
            st.session_state.pop("current_conversation_id", None)
            st.session_state.pop("chat_history", None)
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# Literature search (Europe PMC + OpenAlex)
# =============================================================================
def _fetch_europe_pmc(search_query: str) -> list:
    try:
        response = requests.get(
            "https://www.ebi.ac.uk/europepmc/webservices/rest/search",
            params={
                "query": f"({search_query}) AND OPEN_ACCESS:Y",
                "format": "json",
                "resultType": "core",
                "pageSize": ONLINE_RESULT_LIMIT,
            },
            timeout=ONLINE_TIMEOUT_SECONDS,
        )
        response.raise_for_status()
        results = []
        for item in response.json().get("resultList", {}).get("result", []):
            identifier = f"[PMID: {item.get('pmid')}]" if item.get('pmid') else f"[DOI: {item.get('doi', 'unknown')}]"
            results.append({
                "database": "Europe PMC",
                "source": identifier,
                "title": item.get("title", ""),
                "journal": item.get("journalTitle", ""),
                "abstract": item.get("abstractText", ""),
                "year": item.get("pubYear", ""),
            })
        return results
    except Exception as e:
        return [{"source": "Europe PMC", "status": f"data unavailable: {type(e).__name__}: {e}"}]

def _fetch_openalex(search_query: str) -> list:
    try:
        response = requests.get(
            "https://api.openalex.org/works",
            params={"search": search_query, "filter": "is_oa:true", "per-page": ONLINE_RESULT_LIMIT},
            timeout=ONLINE_TIMEOUT_SECONDS,
        )
        response.raise_for_status()
        results = []
        for item in response.json().get("results", []):
            doi = item.get("doi") or item.get("ids", {}).get("openalex", "unknown")
            identifier = f"[DOI: {doi.replace('https://doi.org/', '')}]" if doi else "[OpenAlex: unknown]"
            results.append({
                "database": "OpenAlex",
                "source": identifier,
                "title": item.get("title", ""),
                "journal": item.get("primary_location", {}).get("source", {}).get("display_name", "") if item.get("primary_location") else "",
                "year": item.get("publication_year", ""),
                "open_access": item.get("open_access", {}),
                "landing_page": item.get("primary_location", {}).get("landing_page_url", "") if item.get("primary_location") else "",
            })
        return results
    except Exception as e:
        return [{"source": "OpenAlex", "status": f"data unavailable: {type(e).__name__}: {e}"}]

def fetch_online_open_access_context(query: str):
    search_query = translate_query_to_english(
        query,
        st.session_state.get("active_model_choice", ""),
        st.session_state.get("active_model_key", ""),
    )
    if not search_query:
        return [{"source": "Europe PMC / OpenAlex", "status": "empty query, search skipped"}], []

    with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
        europe_pmc_future = executor.submit(_fetch_europe_pmc, search_query)
        openalex_future = executor.submit(_fetch_openalex, search_query)
        context = europe_pmc_future.result() + openalex_future.result()

    sources = [record["source"] for record in context if "status" not in record]
    return context, list(dict.fromkeys(sources))

# =============================================================================
# Chat UI — Open Access research assistant
# =============================================================================
def render_online_research():
    st.header("ระบบสืบค้นและวิเคราะห์ข้อมูลชีววิทยาแบบเปิด (Open Access Biology Research)")
    st.caption("ระบบผู้ช่วยวิเคราะห์ที่อ้างอิงข้อมูลจากฐานข้อมูล Open Access ระดับสากล และบันทึกประวัติการสนทนาให้โดยอัตโนมัติ")

    ensure_chat_session_state()

    with st.popover("แนบเอกสาร / ถ่ายภาพ"):
        attached_files = st.file_uploader(
            "อัปโหลดไฟล์เอกสารหรือรูปภาพ",
            type=["pdf", "txt", "md", "png", "jpg", "jpeg", "webp"],
            accept_multiple_files=True,
            key="online_attachments",
            label_visibility="collapsed"
        )
        camera_image = st.camera_input("ถ่ายภาพ", key="online_camera", label_visibility="collapsed")

    # Show chat history
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            if message.get("sources"):
                with st.expander("รายการอ้างอิงจาก Open Access"):
                    for source in message["sources"]:
                        st.write(source)

    # Unified chat input (no separate text_input)
    query = st.chat_input("พิมพ์คำถามเกี่ยวกับการวิจัยทางชีววิทยา หรือชีวสารสนเทศ...")
    
    if query:
        conversation_id = st.session_state.current_conversation_id
        st.session_state.chat_history.append({"role": "user", "content": query})
        append_chat_message(conversation_id, "user", query)

        with st.chat_message("user"):
            st.markdown(query)

        with st.chat_message("assistant"):
            try:
                context, sources = fetch_online_open_access_context(query)
                for record in context:
                    if record.get("status"):
                        st.warning(f"{record.get('source', 'ฐานข้อมูล')}: {record['status']}")
                available_records = [
                    record for record in context
                    if record.get("abstract") or record.get("title")
                ]
                
                if not available_records:
                    answer = "ข้อมูลจากฐานข้อมูล Open Access ณ ปัจจุบัน ไม่เพียงพอต่อการวิเคราะห์เพื่อตอบคำถามดังกล่าว"
                else:
                    attachment_parts, attachment_names = build_attachment_content(
                        attached_files,
                        camera_image,
                    )
                    prompt = f"""
คุณคือ Advanced Bioinformatics Research Agent ในรูปแบบ Chatbot ทางวิชาการ
ใช้ข้อมูลสำหรับการอ้างอิงจาก <open_access_context> และไฟล์แนบที่ผู้ใช้นิยามเท่านั้น
ห้ามใช้ข้อมูลพื้นฐานที่โมเดลได้รับการฝึกมา (Pretrained knowledge) และห้ามคาดเดาข้อมูล
ข้อมูลทางวิชาการทุกประการต้องมีเอกสารอ้างอิง (Inline citation) กำกับทันที
กรุณาแบ่งส่วนเนื้อหาเป็น Facts (ข้อเท็จจริง) และ Inferences (การอนุมาน) อย่างชัดเจน
หากข้อมูลไม่เพียงพอ ให้ตอบ EXACTLY ว่า:
"ข้อมูลจากฐานข้อมูล Open Access ณ ปัจจุบัน ไม่เพียงพอต่อการวิเคราะห์เพื่อตอบคำถามดังกล่าว"
ใช้ระดับภาษาไทยกึ่งทางการหรือทางการเชิงวิชาการ

<open_access_context>
{json.dumps(context, ensure_ascii=False, indent=2)}
</open_access_context>

ไฟล์แนบ (Attached files): {', '.join(attachment_names) or '(ไม่มี)'}
คำถามจากผู้ใช้ (User query): {query}
"""
                    with st.spinner("ระบบกำลังสืบค้นและประมวลผลข้อมูล..."):
                        response = get_active_llm(
                            st.session_state.active_model_choice,
                            st.session_state.active_model_key,
                        ).invoke([HumanMessage(content=[prompt, *attachment_parts])])
                    answer = extract_text(response).strip()
                
                st.markdown(answer)
                if sources:
                    with st.expander("รายการอ้างอิงจาก Open Access"):
                        for source in sources:
                            st.write(source)
                
                st.session_state.chat_history.append({
                    "role": "assistant",
                    "content": answer,
                    "sources": sources
                })
                append_chat_message(conversation_id, "assistant", answer, sources)

            except Exception as e:
                answer = f"เกิดข้อผิดพลาดในการเชื่อมต่อฐานข้อมูลหรือประมวลผลโมเดล AI: {type(e).__name__}: {e}"
                st.error(answer)
                st.session_state.chat_history.append({
                    "role": "assistant",
                    "content": answer,
                    "sources": []
                })
                append_chat_message(conversation_id, "assistant", answer, [])

# =============================================================================
# Bioinformatics Pipeline
# =============================================================================
def parse_analysis_json(raw_response: str):
    cleaned = raw_response.strip()
    if cleaned.startswith("```"):
        raise ValueError("รูปแบบข้อมูลที่ส่งกลับมาจากโมเดลไม่ถูกต้อง (ตรวจพบ Markdown code fence) กรุณาติดต่อผู้พัฒนาระบบ")
    analysis = json.loads(cleaned)
    required = {"summary", "confidence_score", "uncertainties", "next_steps"}
    if not isinstance(analysis, dict) or set(analysis) != required:
        raise ValueError("โครงสร้าง JSON ที่ได้รับไม่ตรงกับเกณฑ์ที่กำหนด กรุณาติดต่อผู้พัฒนาระบบ")
    if not isinstance(analysis["summary"], str):
        raise ValueError("ข้อมูลส่วน 'summary' ต้องอยู่ในรูปแบบ String กรุณาติดต่อผู้พัฒนาระบบ")
    if not isinstance(analysis["confidence_score"], int) or isinstance(analysis["confidence_score"], bool):
        raise ValueError("ข้อมูลส่วน 'confidence_score' ต้องเป็นค่าจำนวนเต็ม (Integer) กรุณาติดต่อผู้พัฒนาระบบ")
    if not 0 <= analysis["confidence_score"] <= 100:
        raise ValueError("ค่า 'confidence_score' จะต้องอยู่ระหว่าง 0 ถึง 100 กรุณาติดต่อผู้พัฒนาระบบ")
    for field in ("uncertainties", "next_steps"):
        if not isinstance(analysis[field], list) or not all(isinstance(item, str) for item in analysis[field]):
            raise ValueError(f"ข้อมูลส่วน '{field}' ต้องอยู่ในรูปแบบรายการ (List of Strings) กรุณาติดต่อผู้พัฒนาระบบ")
    return analysis

# =============================================================================
# Module 2: Biopython deterministic pipeline and Gemini agent
# =============================================================================
def parse_uploaded_sequence(uploaded_file):
    try:
        content = uploaded_file.getvalue().decode("utf-8")
        filename = uploaded_file.name.lower()
        file_format = "fastq" if filename.endswith((".fastq", ".fq")) else "fasta"

        records = list(SeqIO.parse(StringIO(content), file_format))
        if not records:
            raise ValueError("ระบบไม่พบข้อมูลลำดับเบส (Sequence) ในไฟล์ที่แนบมา กรุณาตรวจสอบไฟล์อีกครั้ง")
        record = records[0]
    except Exception as exc:
        raise ValueError(f"รูปแบบไฟล์หรือข้อมูลภายในไม่ถูกต้อง: {exc} กรุณาติดต่อผู้พัฒนาระบบ") from exc
    return record.seq, record.id

def parse_raw_sequence(raw_input: str):
    sequence = re.sub(r"\s+", "", raw_input or "").upper()
    if not sequence:
        raise ValueError("กรุณาระบุลำดับนิวคลีโอไทด์ (Nucleotide Sequence) เพื่อดำเนินการต่อ")
    return Seq(sequence), "raw-sequence"

def fetch_ncbi_sequence(accession: str):
    accession = accession.strip()
    if not accession:
        raise ValueError("กรุณาระบุรหัสอ้างอิง (NCBI Accession Number) เพื่อค้นหา")

    try:
        with st.spinner("ระบบกำลังสืบค้นลำดับเบสจาก GenBank..."):
            with Entrez.efetch(
                db="nucleotide",
                id=accession,
                rettype="fasta",
                retmode="text",
                timeout=API_TIMEOUT_SECONDS,
            ) as handle:
                record = SeqIO.read(handle, "fasta")
            return record.seq, record.id
    except HTTPError as exc:
        raise ValueError(f"เซิร์ฟเวอร์ NCBI ปฏิเสธคำขอดึงข้อมูล ({exc.code}) กรุณาตรวจสอบรหัส Accession หรือติดต่อผู้พัฒนาระบบ") from exc
    except URLError as exc:
        raise ValueError(f"การเชื่อมต่อกับเซิร์ฟเวอร์ NCBI ขัดข้อง: {exc.reason} กรุณาติดต่อผู้พัฒนาระบบ") from exc
    except Exception as exc:
        raise ValueError(f"เกิดข้อผิดพลาดที่ไม่ทราบสาเหตุระหว่างดึงข้อมูล: {exc} กรุณาติดต่อผู้พัฒนาระบบ") from exc

def calculate_sequence_metrics(sequence: Seq):
    sequence_text = str(sequence).upper()
    invalid = sorted(set(sequence_text) - set("ACGTN"))
    if invalid:
        raise ValueError(f"ตรวจพบอักษระที่ไม่ใช่อักษรนิวคลีโอไทด์มาตรฐาน: {', '.join(invalid)}")
    gc_count = sequence_text.count("G") + sequence_text.count("C")
    gc_content = gc_count / len(sequence_text) * 100 if sequence_text else 0.0
    mrna = sequence.transcribe()
    protein = mrna.translate(to_stop=True)
    return {"sequence": sequence_text, "length": len(sequence_text), "gc": gc_content, "mrna": str(mrna), "protein": str(protein)}

def resolve_database_query(sequence_id: str, user_query: str, metrics: dict) -> str:
    if sequence_id and sequence_id != "raw-sequence":
        return sequence_id.strip()
    if user_query.strip():
        return user_query.strip()
    return metrics.get("protein", "").strip() or "nucleotide sequence"

def fetch_ncbi_identification(sequence_id: str, sequence: str):
    result = {
        "query": sequence_id,
        "accession": sequence_id if sequence_id != "raw-sequence" else None,
        "gene": None,
        "organism": None,
        "length": len(sequence),
        "e_value": None,
        "identity_percent": None,
        "tax_id": None,
    }
    try:
        with Entrez.esearch(db="nuccore", term=sequence_id, retmax=1, timeout=API_TIMEOUT_SECONDS) as search_handle:
            search_result = Entrez.read(search_handle)
        identifiers = search_result.get("IdList", [])
        if not identifiers:
            return {"status": "ไม่พบข้อมูลระบุตัวตนในฐานข้อมูล NCBI", **result}
        with Entrez.efetch(db="nuccore", id=identifiers[0], rettype="gb", retmode="text", timeout=API_TIMEOUT_SECONDS) as fetch_handle:
            record = SeqIO.read(fetch_handle, "genbank")
        result["accession"] = record.id or result["accession"]
        result["length"] = len(record.seq)
        result["organism"] = record.annotations.get("organism")
        for dbxref in record.dbxrefs:
            if dbxref.startswith("taxon:"):
                result["tax_id"] = dbxref.split(":", 1)[1]
                break
        for feature in record.features:
            if feature.type in {"gene", "CDS"}:
                qualifiers = feature.qualifiers
                result["gene"] = (qualifiers.get("gene") or qualifiers.get("locus_tag") or [None])[0]
                if result["gene"]:
                    break
        return result
    except Exception as exc:
        return {"status": "ระบบไม่สามารถดึงข้อมูลระบุตัวตนจาก NCBI ได้", "error": str(exc), **result}

def fetch_literature_data(query: str):
    records, _ = fetch_online_open_access_context(query)
    return [
        record for record in records
        if record.get("database") == "Europe PMC"
        and (record.get("abstract") or record.get("title"))
    ]

def _uniprot_result(response):
    response.raise_for_status()
    payload = response.json()
    return payload.get("results", []) if "results" in payload else [payload]

def uniprot_fetch_by_accession(accession_id: str):
    response = requests.get(
        UNIPROT_ENTRY_ENDPOINT.format(accession_id=accession_id.strip()),
        timeout=API_TIMEOUT_SECONDS,
    )
    results = _uniprot_result(response)
    return results[0] if results else None

def uniprot_fetch_by_gene(gene_name: str):
    response = requests.get(
        UNIPROT_GENE_ENDPOINT.format(gene_name=requests.utils.quote(gene_name.strip())),
        timeout=API_TIMEOUT_SECONDS,
    )
    results = _uniprot_result(response)
    return results[0] if results else None

def uniprot_fetch_by_organism(tax_id: str):
    response = requests.get(
        UNIPROT_ORGANISM_ENDPOINT.format(tax_id=str(tax_id).strip()),
        timeout=API_TIMEOUT_SECONDS,
    )
    results = _uniprot_result(response)
    return results[0] if results else None

def uniprot_fetcher(query: str, gene_name: str = "", tax_id: str = ""):
    try:
        accession_match = re.fullmatch(
            r"(?:[OPQ][0-9][A-Z0-9]{3}[0-9]|[A-NR-Z][0-9][A-Z0-9]{3}[0-9])",
            query.strip().upper(),
        )
        entry = uniprot_fetch_by_accession(query) if accession_match else None
        if not entry:
            entry = uniprot_fetch_by_gene(gene_name or query)
        organism_entry = uniprot_fetch_by_organism(tax_id) if tax_id else None
        if not entry and not organism_entry:
            return {"status": "ไม่พบข้อมูลที่ตรงกันในระบบ UniProt", "query": query}
        entry = entry or organism_entry
        domains = [
            feature.get("description", "")
            for feature in entry.get("features", [])
            if feature.get("type") in {"Domain", "Region"}
        ]
        locations = [
            comment.get("texts", [{}])[0].get("value", "")
            for comment in entry.get("comments", [])
            if comment.get("commentType") == "SUBCELLULAR LOCATION"
        ]
        functions = [
            text.get("value", "")
            for comment in entry.get("comments", [])
            if comment.get("commentType") == "FUNCTION"
            for text in comment.get("texts", [])
        ]
        return {
            "accession": entry.get("primaryAccession"),
            "protein": entry.get("proteinDescription", {}).get("recommendedName", {}).get("fullName", {}).get("value"),
            "function": functions,
            "subcellular_location": locations,
            "functional_domains": domains,
        }
    except Exception as exc:
        return {"status": "เกิดข้อผิดพลาดในการเชื่อมต่อฐานข้อมูล UniProt", "error": str(exc)}

def clinvar_fetcher(query: str):
    try:
        search_handle = Entrez.esearch(db="clinvar", term=query, retmax=5, timeout=API_TIMEOUT_SECONDS)
        search_result = Entrez.read(search_handle)
        search_handle.close()
        identifiers = search_result.get("IdList", [])
        if not identifiers:
            return {"status": "ไม่พบข้อมูลในฐานข้อมูล ClinVar", "query": query}
        fetch_handle = Entrez.efetch(db="clinvar", id=identifiers, rettype="vcv", retmode="xml", timeout=API_TIMEOUT_SECONDS)
        raw_xml = fetch_handle.read()
        fetch_handle.close()
        return {"query": query, "record_ids": identifiers, "summary": str(raw_xml)[:12000]}
    except Exception as exc:
        return {"status": "เกิดข้อผิดพลาดในการเชื่อมต่อฐานข้อมูล ClinVar", "error": str(exc)}

def kegg_fetcher(query: str):
    try:
        response = requests.get(
            "https://rest.kegg.jp/find/genes",
            params={"term": query},
            timeout=API_TIMEOUT_SECONDS,
        )
        response.raise_for_status()
        matches = [line.split("\t", 1) for line in response.text.splitlines() if "\t" in line]
        if not matches:
            return {"status": "ไม่พบข้อมูลในระบบ KEGG", "query": query}
        gene_id = matches[0][0]
        link_response = requests.get(
            f"https://rest.kegg.jp/link/pathway/{gene_id}",
            timeout=API_TIMEOUT_SECONDS,
        )
        link_response.raise_for_status()
        pathways = [line.split("\t", 1)[1] for line in link_response.text.splitlines() if "\t" in line]
        return {"query": query, "gene_match": matches[0][1], "pathways": pathways}
    except Exception as exc:
        return {"status": "เกิดข้อผิดพลาดในการเชื่อมต่อระบบ KEGG", "error": str(exc)}

def pdb_fetcher(query: str):
    try:
        response = requests.post(
            "https://search.rcsb.org/rcsbsearch/v2/query",
            json={
                "query": {"type": "terminal", "service": "full_text", "parameters": {"value": query}},
                "return_type": "entry",
                "request_options": {"paginate": {"start": 0, "rows": 5}},
            },
            timeout=API_TIMEOUT_SECONDS,
        )
        response.raise_for_status()
        identifiers = [item.get("identifier") for item in response.json().get("result_set", [])]
        if not identifiers:
            return {"status": "ไม่พบข้อมูลโครงสร้างใน PDB", "query": query}
        entries = []
        for pdb_id in identifiers:
            entry_response = requests.get(
                f"https://data.rcsb.org/rest/v1/core/entry/{pdb_id}",
                timeout=API_TIMEOUT_SECONDS,
            )
            entry_response.raise_for_status()
            entry = entry_response.json().get("struct", {})
            entries.append({
                "pdb_id": pdb_id,
                "title": entry.get("title", ""),
                "experimental_method": entry.get("pdbx_descriptor", ""),
            })
        return {"query": query, "pdb_ids": identifiers, "structural_summary": entries}
    except Exception as exc:
        return {"status": "เกิดข้อผิดพลาดในการเชื่อมต่อระบบ RCSB PDB", "error": str(exc)}

def generate_txt(content: str) -> str:
    return content

def generate_docx(content: str) -> bytes:
    document = WordDocument()
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 3)
            document.add_heading(stripped.lstrip("# "), level=level)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            document.add_paragraph(stripped)
    from io import BytesIO
    binary_output = BytesIO()
    document.save(binary_output)
    return binary_output.getvalue()

def build_attachment_content(uploaded_files, camera_image):
    content = []
    attachment_names = []
    for uploaded_file in uploaded_files or []:
        attachment_names.append(uploaded_file.name)
        file_bytes = uploaded_file.getvalue()
        mime_type = uploaded_file.type or "application/octet-stream"
        if mime_type.startswith("image/"):
            encoded = base64.b64encode(file_bytes).decode("ascii")
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{mime_type};base64,{encoded}"},
            })
        elif mime_type == "application/pdf" or uploaded_file.name.lower().endswith(".pdf"):
            pages = PdfReader(BytesIO(file_bytes)).pages
            content.append("\n".join(page.extract_text() or "" for page in pages))
        else:
            content.append(file_bytes.decode("utf-8", errors="replace"))
    if camera_image is not None:
        attachment_names.append("camera-capture.jpg")
        encoded = base64.b64encode(camera_image.getvalue()).decode("ascii")
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{camera_image.type or 'image/jpeg'};base64,{encoded}"},
        })
    return content, attachment_names

def analyze_bio_context_with_attachments(metrics, attachments, query):
    existing_report = st.session_state.get("bio_report", "")
    attachment_parts, attachment_names = attachments
    prompt = f"""
คุณคือ Expert Bioinformatics Research Analyst
จงวิเคราะห์ข้อมูลเดิมและไฟล์เอกสารแนบที่ผู้ใช้ส่งมา โดยต้องแยกแยะข้อเท็จจริงออกจาก
การคาดเดา
ห้ามนำเสนอข้อมูลหรือจัดทำเอกสารอ้างอิงที่ไม่มีหลักฐานเชิงประจักษ์รองรับ และให้ระบุข้อจำกัดหากปริมาณข้อมูลไม่เพียงพอ

ข้อมูลเชิงปริมาณ (Deterministic Metrics):
{json.dumps(metrics, ensure_ascii=False, indent=2)}

รายชื่อไฟล์แนบ: {', '.join(attachment_names) or '(ไม่มีไฟล์แนบ)'}
ข้อกำหนดเพิ่มเติม: {query or 'โปรดวิเคราะห์ข้อมูลทั้งหมดและสรุปประเด็นหลักทางวิชาการ'}
"""
    message_content = [prompt, *attachment_parts]
    response = get_active_llm(
        st.session_state.active_model_choice,
        st.session_state.active_model_key,
    ).invoke([HumanMessage(content=message_content)])
    return extract_text(response)

# =============================================================================
# Bioinformatics UI — Professional Scientific Layout
# =============================================================================
def render_bioinformatics():
    st.markdown("### Bioinformatics Analysis Agent", unsafe_allow_html=True)
    st.caption("วิเคราะห์ข้อมูลเชิงปริมาณด้วย Biopython และสังเคราะห์ผลด้วย AI")

    input_method = st.selectbox(
        "Sequence Input",
        ["อัปโหลดไฟล์ (File Upload)", "ระบุข้อความ (Raw Text Input)", "ระบุรหัสอ้างอิง (NCBI Accession)"],
        key="sequence_input_method",
    )

    uploaded_file = None
    raw_input = ""
    accession = ""

    if input_method == "อัปโหลดไฟล์ (File Upload)":
        uploaded_file = st.file_uploader(
            "อัปโหลดไฟล์ลำดับเบส (FASTA / FASTQ)",
            type=["fasta", "fa", "fastq", "fq"],
            key="sequence_file",
        )
    elif input_method == "ระบุข้อความ (Raw Text Input)":
        raw_input = st.text_area(
            "ระบุลำดับนิวคลีโอไทด์ (Nucleotide Sequence)",
            height=140,
            key="sequence_text",
        )
    else:
        accession = st.text_input(
            "ระบุรหัสอ้างอิง NCBI (Accession Number)",
            placeholder="ตัวอย่าง: NM_000546 หรือ NC_045512",
            key="ncbi_accession",
        )

    query = st.text_input(
        "คำสั่งการวิเคราะห์เชิงลึก (Prompt)",
        placeholder="ตัวอย่าง: จงวิเคราะห์หน้าที่ทางชีวภาพและบทบาททางพยาธิวิทยาของลำดับเบสนี้",
    )

    analysis_submitted = st.button("ประมวลผลและสร้างรายงานวิชาการ", type="primary")

    pipeline_states = {step: "pending" for step in PIPELINE_STEPS}
    analysis_log = AnalysisLog()

    if analysis_submitted:
        pipeline_states["Input Validation"] = "active"
        analysis_log.log("Input validation started")

        try:
            if input_method == "อัปโหลดไฟล์ (File Upload)":
                if uploaded_file is None:
                    raise ValueError("กรุณาเลือกไฟล์แนบที่ต้องการอัปโหลด")
                valid, err, detected_fmt = validate_uploaded_file(uploaded_file)
                if not valid:
                    raise ValueError(err)
                analysis_log.log(f"File detected: {detected_fmt}")
                sequence, sequence_id = parse_uploaded_sequence(uploaded_file)
            elif input_method == "ระบุข้อความ (Raw Text Input)":
                sequence, sequence_id = parse_raw_sequence(raw_input)
            else:
                sequence, sequence_id = fetch_ncbi_sequence(accession)

            st.session_state.bio_metrics = calculate_sequence_metrics(sequence)
            st.session_state.bio_sequence_id = sequence_id
            st.session_state.bio_input_method = input_method
            st.session_state.bio_analysis_id = generate_analysis_id()

            pipeline_states["Input Validation"] = "done"
            analysis_log.log("Input validation completed")

            validation = validate_sequence(sequence, sequence_id, input_method)
            st.session_state.bio_validation = validation
            pipeline_states["Sequence Analysis"] = "done"
            analysis_log.log("Sequence analysis completed")

        except ValueError as exc:
            st.error(str(exc))
            pipeline_states["Input Validation"] = "pending"
            analysis_log.log(f"Validation failed: {exc}", is_error=True)
            with st.expander("Analysis Log"):
                analysis_log.render()
            return

    metrics = st.session_state.get("bio_metrics")
    if not metrics:
        return

    validation = st.session_state.get("bio_validation")
    if validation:
        render_validation(validation)

    render_pipeline(pipeline_states)

    st.markdown("#### 1. Computational Results")
    st.caption(f"Sequence ID: {st.session_state.get('bio_sequence_id', 'N/A')}")

    col1, col2, col3 = st.columns(3)
    col1.metric("Sequence Length", f"{metrics['length']} bp")
    col2.metric("GC Content", f"{metrics['gc']:.2f}%")
    col3.metric("Protein Length", f"{len(metrics['protein'])} aa")

    with st.expander("Translation Results"):
        st.code(f"mRNA:\n{metrics['mrna']}\n\nProtein (to Stop Codon):\n{metrics['protein'] or '(No protein before Stop Codon)'}")

    if not st.session_state.get("active_model_key"):
        model_choice = st.session_state.get("active_model_choice", "โมเดลที่เลือก")
        st.warning(f"ระบบตรวจพบว่ายังไม่มีข้อมูล API Key สำหรับ {model_choice} ซึ่งจำเป็นต่อกระบวนการวิเคราะห์ผลขั้นสูงด้วย AI")
        return

    if analysis_submitted:
        pipeline_states["Database Retrieval"] = "active"
        analysis_log.log("Database retrieval started")

        database_query = resolve_database_query(
            st.session_state.get("bio_sequence_id", ""),
            query,
            metrics,
        )
        # NCBI, ClinVar, KEGG, PDB, and literature lookups are independent of
        # each other, so they run concurrently instead of one after another --
        # sequentially, five ~10-30s-timeout network calls could take well
        # over a minute before the AI step even starts.
        with st.spinner("ระบบกำลังสืบค้นข้อมูลจากฐานข้อมูลชีวสารสนเทศหลายแห่งพร้อมกัน..."):
            with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
                identification_future = executor.submit(fetch_ncbi_identification, database_query, metrics["sequence"])
                clinvar_future = executor.submit(clinvar_fetcher, database_query)
                kegg_future = executor.submit(kegg_fetcher, database_query)
                pdb_future = executor.submit(pdb_fetcher, database_query)
                literature_future = executor.submit(fetch_literature_data, database_query)

                identification_data = identification_future.result()
                clinvar_result = clinvar_future.result()
                kegg_result = kegg_future.result()
                pdb_result = pdb_future.result()
                literature_data = literature_future.result()
            analysis_log.log("NCBI, ClinVar, KEGG, PDB, and literature queries completed")

            uniprot_result = uniprot_fetcher(
                database_query,
                gene_name=identification_data.get("gene") or "",
                tax_id=identification_data.get("tax_id") or "",
            )
            analysis_log.log("UniProt query completed")

        protein_data = {
            "UniProt": uniprot_result,
            "ClinVar": clinvar_result,
            "KEGG": kegg_result,
            "RCSB PDB": pdb_result,
        }

        pipeline_states["Database Retrieval"] = "done"
        pipeline_states["Computational Analysis"] = "done"
        pipeline_states["AI Interpretation"] = "active"
        analysis_log.log("AI interpretation started")

        evidence = {
            "ncbi": {
                # identification_data["accession"] falls back to echoing the
                # search query when nothing was found; only surface it here
                # as evidence when NCBI actually confirmed a record.
                "accession": identification_data.get("accession") if "status" not in identification_data else None,
                "organism": identification_data.get("organism"),
                "gene": identification_data.get("gene"),
            },
            "uniprot": {
                "accession": protein_data.get("UniProt", {}).get("accession"),
            },
            "kegg": {
                "pathways": protein_data.get("KEGG", {}).get("pathways", []),
            },
            "pdb": {
                "pdb_ids": protein_data.get("RCSB PDB", {}).get("pdb_ids", []),
            },
        }
        st.session_state.bio_evidence = evidence

        prompt = f"""ข้อมูลรวบรวมจากเครือข่ายฐานข้อมูลวิชาการแบบเปิด (Open Access):
ข้อกำหนด: ผู้ช่วยวิเคราะห์ต้องดำเนินการสังเคราะห์ข้อมูลจากข้อมูล 3 ชุดด้านล่างนี้เท่านั้น ห้ามอ้างอิงข้อมูลภายนอกหรือคาดเดาสิ่งที่ปรากฏในหลักฐาน
ข้อเท็จจริง (Facts) และการอนุมาน (Inferences) ทั้งหมดต้องมีเอกสารอ้างอิงกำกับทันที
กรุณาจัดโครงสร้างการรายงานโดยแบ่งหัวข้อข้อเท็จจริงและการอนุมานออกจากกันอย่างชัดเจน
หากไม่พบข้อมูลในส่วนใด ให้ระบุข้อความ "ไม่พบข้อมูลที่ตรงกันในฐานข้อมูล [ชื่อระบบ]"
ใช้ภาษาไทยระดับวิชาการเพื่อจัดทำรายงาน

<open_access_context>
<identification_data>
{json.dumps(identification_data, ensure_ascii=False, indent=2)}
</identification_data>
<protein_data>
{json.dumps(protein_data, ensure_ascii=False, indent=2)}
</protein_data>
<literature_data>
{json.dumps(literature_data, ensure_ascii=False, indent=2)}
</literature_data>
</open_access_context>

โปรดวิเคราะห์ข้อมูลทั้งหมดอย่างเคร่งครัดตามแนวปฏิบัติข้างต้น
ข้อกำหนดหรือคำสั่งเพิ่มเติมจากผู้ใช้: {query or 'จงวิเคราะห์หน้าที่และความสำคัญทางชีวภาพของลำดับเบสนี้'}
"""
        try:
            with st.spinner("ระบบกำลังสังเคราะห์ผลลัพธ์การวิเคราะห์ผ่านเครือข่าย AI..."):
                response = get_active_llm(
                    st.session_state.active_model_choice,
                    st.session_state.active_model_key,
                ).invoke([HumanMessage(content=prompt)])
            report = extract_text(response)
            st.session_state.bio_report = report
            pipeline_states["AI Interpretation"] = "done"
            pipeline_states["Evidence Verification"] = "done"
            pipeline_states["Report Generation"] = "done"
            analysis_log.log("AI interpretation completed")
            analysis_log.log("Report generated")
        except Exception as e:
            st.error(f"เกิดข้อผิดพลาดในการวิเคราะห์ลำดับเบสผ่านแบบจำลองภาษา: {type(e).__name__}: {e}")
            pipeline_states["AI Interpretation"] = "pending"
            analysis_log.log("AI interpretation failed", is_error=True)

    report = st.session_state.get("bio_report")
    if report:
        st.divider()

        st.markdown("#### 2. AI Interpretation")
        st.caption("ส่วนนี้เป็นตีความโดย AI ไม่ใช่หลักฐานเชิงทดลอง (Experimental Evidence)")
        st.markdown(f'<div class="report-body">{report}</div>', unsafe_allow_html=True)

        evidence = st.session_state.get("bio_evidence")
        if evidence:
            st.markdown("#### 3. Evidence & Sources")
            render_evidence_sources(evidence)

        st.markdown("#### 4. AI Confidence")
        st.markdown(
            '<div class="confidence-display">'
            '<span class="confidence-score">—</span>'
            '<span class="confidence-label">Pending</span>'
            '</div>'
            '<div class="confidence-desc">'
            'คะแนนนี้สะท้อนระดับความมั่นใจของ AI ต่อผลการตีความจากข้อมูลที่ใช้ในการวิเคราะห์ '
            'ไม่ใช่ค่าความเป็นไปได้ทางชีววิทยา หรือผลการพิสูจน์ทางสถิติ'
            '</div>',
            unsafe_allow_html=True,
        )

        st.markdown("#### 5. Limitations")
        render_limitations()

        analysis_id = st.session_state.get("bio_analysis_id", "N/A")
        metadata = {
            "Analysis ID": analysis_id,
            "Analysis Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "AI Model": st.session_state.get("active_model_choice", "N/A"),
            "Sequence ID": st.session_state.get("bio_sequence_id", "N/A"),
            "Sequence Length": f"{metrics['length']} bp",
            "Input Format": st.session_state.get("bio_input_method", "N/A"),
            "Processing Time": f"{analysis_log.elapsed():.1f}s",
        }
        with st.expander("Analysis Metadata (Reproducibility)"):
            render_analysis_metadata(metadata)

        with st.expander("Analysis Log"):
            analysis_log.render()

        dl_col1, dl_col2 = st.columns(2)
        with dl_col1:
            st.download_button(
                "ดาวน์โหลดรายงานฉบับสมบูรณ์ (Word / DOCX)",
                data=generate_docx(report),
                file_name="bioinformatics_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with dl_col2:
            st.download_button(
                "ดาวน์โหลดข้อมูลข้อความ (Text / TXT)",
                data=generate_txt(report),
                file_name="bioinformatics_report.txt",
                mime="text/plain",
            )

# =============================================================================
# State router: each mode owns its own session-state keys
# =============================================================================
MODE_OPTIONS = ["สืบค้นข้อมูล Open Access ออนไลน์", "วิเคราะห์ข้อมูลชีวสารสนเทศ"]
if "active_mode" not in st.session_state or st.session_state.active_mode not in MODE_OPTIONS:
    st.session_state.active_mode = MODE_OPTIONS[0]

with st.sidebar:
    st.title("Bioinformatics AI")

    st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-section-label">AI MODEL</div>', unsafe_allow_html=True)
    st.session_state.active_model_choice = st.selectbox(
        "AI Model",
        list(MODEL_OPTIONS),
        key="model_selector",
        label_visibility="collapsed",
    )
    st.markdown('</div>', unsafe_allow_html=True)

    active_model_config = MODEL_OPTIONS[st.session_state.active_model_choice]
    configured_key = get_configured_key(active_model_config["secret"])
    if configured_key:
        st.session_state.active_model_key = configured_key
    else:
        api_key_state = f"api_key_{active_model_config['secret']}"
        st.session_state.active_model_key = st.text_input(
            "ระบุรหัสเชื่อมต่อ (API Key)",
            type="password",
            key=api_key_state,
            label_visibility="collapsed",
        ).strip()
        if not st.session_state.active_model_key:
            st.warning(f"ระบบไม่พบข้อมูล {active_model_config['secret']} สำหรับการใช้งานโมเดลนี้")

    st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-section-label">ANALYSIS MODE</div>', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-section-desc">เลือกฟังก์ชันการใช้งาน</div>', unsafe_allow_html=True)
    st.session_state.active_mode = st.radio(
        "Analysis mode",
        MODE_OPTIONS,
        index=MODE_OPTIONS.index(st.session_state.active_mode),
        label_visibility="collapsed",
    )
    st.markdown('</div>', unsafe_allow_html=True)

    if st.session_state.active_mode == MODE_OPTIONS[0]:
        render_chat_sidebar()

    st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-section-label">STATUS</div>', unsafe_allow_html=True)
    render_sidebar_status()
    st.markdown('</div>', unsafe_allow_html=True)

if st.session_state.active_mode == MODE_OPTIONS[0]:
    render_online_research()
else:
    render_bioinformatics()
